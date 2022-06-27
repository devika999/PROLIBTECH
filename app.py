import base64
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

import pandas as pd
from flask import Flask,render_template,request,session,redirect,jsonify
# import pandas as pd
from DBConnection import Db
import datetime
import time
app = Flask(__name__)
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
import aspose.words as aw
import nltk
nltk.download('stopwords')
stop_words = set(stopwords.words('english'))
staticpath = "C:\\Users\\User\\PycharmProjects\\project_management12\\static\\"
app.secret_key="ok"


gl_list=[".","-","(",")","_","'",":",",","[","]","@","%","?","/","!","“",":",";","``","a","b","c","d","e","f","g","h",
         "i","j","k","l","m","n","o","p","q","r","s","u","v","w","x","y","z","1","2","3","4","5","6","7","8","9","10","0","’",
         "...................","i.e","3.4.3","3.3","``","will","only","pty","ttd.","of","no","reg","fulfillmentof","award","degree","of","and","dist","hereby","declare",
         "sub-","mitted","Insti-","Date","the","place","camscanner","this","certify","carried","year","Guide","held","Mr.","de-","gav","key","accuracy","did","new",
         "due","dew","not","tea","note","come-on","concerning","concerninger","concerningest","consequently","considering","eg","eight",
         "your","yours","yourself","yourselves","you","yond","yonder","yon","ye","yet","zillion","umpteen","usually","username","uponed","upons","uponing","upon",

         "ups", "upping", "upped", "up", "unto", "until", "unless", "unlike", "unliker", "unlikest", "under"
                                                                                                     "underneath",
         "use", "used", "usedest", "rath", "rather", "rathest", "rathe", "re", "relate", "related", "relatively",
         "regarding", "really", "res", "respecting", "respectively", "quite","que","qua","neither","neaths","neath","nethe","nethermost","necessary","necessariest","necessarier","never","nevertheless","nigh","nighest","nigher","nine","noone","nobody",
"ten","tills","till","tilled","tilling","to","towards","toward","towardest","towarder","together","too","thy","thyself","thus","than","that",
"those","thou","though","thous","thouses","thoroughest","thorougher","thorough","thoroughly","thru","thruer","thruest","thro","through","throughout","throughest",
"througher","thine","this","thises","they","thee","the","{","}","[","]"
"nobodies","nowhere","nowheres","no","noes","or","nos","no-one","none","not","notwithstanding","nothings","nothing","nathless","natheless"]
@app.route('/index_page')
def index_page():
    return render_template('searchpr.html')

@app.route('/choosedepartment')
def choosedepartment():
    db = Db()
    res1 = db.select("SELECT * from college")
    print(res1)
    res2 = db.select("SELECT * FROM `department`")
    print(res2)
    return render_template('choosedepartment.html',col=res1,dept=res2)
@app.route('/choosedepartment1')
def choosedepartment1():
    db = Db()
    res1 = db.select("SELECT * from college")
    print(res1)
    res2 = db.select("SELECT * FROM `department`")
    print(res2)
    return render_template('choosedepartment1.html',col=res1,dept=res2)
@app.route('/choosedepartment1_post',methods=["post"])
def choosedepartment1_post():

    dept = request.form['select2']
    session["cdepid"]=dept
    db = Db()
    # res = db.select("SELECT project.*,student.s_name,department.dept_name,branch.branch_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN department ON student.dept_id=department.dept_id INNER JOIN `branch` ON student.branch_id=branch.branch_id WHERE department.dept_id='" + str(dept) + "'")
    return render_template('/searchpp1.html')

@app.route('/searchpp', methods=["post"])
def searchpp():
    session["cdepid"]
    selectp = request.form['select']
    search = request.form['textbox10']
    yr1 = request.form['select2']
    yr2 = request.form['select3']
    db = Db()
    if selectp == "all":
        restt = search.split(" ")
        print(restt)
        lss = []
        ls = []
        pid = []
        qry = "SELECT project.project_id,project.abstract,journal,project.report,githublink,demo,project.project_area,dept_name,project.title,student.s_name,college.c_name,project.p_year,department.dept_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN `ngram` ON `ngram`.`pid`=`project`.`project_id` inner join `department` ON `department`.`dept_id`=`student`.`dept_id` WHERE (department.dept_id='"+str(session["cdepid"])+"') and (ngram.`word` = '" + search + "' or project.`title` = '" + search + "') and  p_year between '"+yr1+"' and '"+yr2+"'"
        print(qry)
        db = Db()
        res = db.select(qry)
        print(len(res), "------")
        if len(res) > 0:
            print(res)
            print(len(res))
            for i in res:
                if i["project_id"] in pid:
                    print("---", i["project_id"])
                    pass
                else:
                    print(pid, "--------------")
                    pid.append(i["project_id"])
                    a = {'title': i['title'], 'abstract': i['abstract'], 'project_id': i['project_id'],
                                                  's_name': i['s_name'], 'project_area': i['project_area'],
                                                  'dept_name': i['dept_name'],
                                                  'c_name': i['c_name'], 'p_year': i['p_year'], 'report': i['report'],
                                                  'journal': i['journal'], 'githublink': i['githublink'], 'demo': i['demo'],
                                                  'status': 'yes'}
                    ls.append(a)
            return render_template('a.html', data=ls)
        else:
            # ss = db.select(" SELECT project.project_id,project.title,student.s_name,college.c_name,project.p_year FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id where project.`title` LIKE '%" + search + "%' or project.`p_year`LIKE '%" + search + "%'")
            ss = db.select(
                " SELECT project.project_id,project.abstract,project.project_area,journal,project.report,githublink,demo,dept_name,project.title,student.s_name,college.c_name,project.p_year FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id inner join `department` ON `department`.`dept_id`=`student`.`dept_id` where (department.dept_id='"+str(session["cdepid"])+"') and  p_year between '"+yr1+"' and '"+yr2+"'")
            print(ss)
            pid = []
            ls = []
            for i in ss:
                p = 0
                for j in restt:
                    if j != '':
                        Q = "SELECT distinct `key_words`.`pid`,key_words.count, project.project_id,project.title,student.s_name,college.c_name,project.p_year FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN `key_words`  ON `project`.`project_id`=`key_words`.`pid` where `key_words`.`pid`='" + str(
                            i["project_id"]) + "' and key_words.`key` ='" + str(j) + "'"
                        pp = db.select(Q)
                        print(Q)
                        if len(pp) > 0:

                            if i["project_id"] in pid:
                                print("---", i["project_id"])
                                pass
                            else:
                                print(pid, "--------------")
                                pid.append(i["project_id"])
                                a = {'title': i['title'], 'abstract': i['abstract'], 'project_id': i['project_id'],
                                                              's_name': i['s_name'], 'project_area': i['project_area'],
                                                              'dept_name': i['dept_name'],
                                                              'c_name': i['c_name'], 'p_year': i['p_year'], 'report': i['report'],
                                                              'journal': i['journal'], 'githublink': i['githublink'], 'demo': i['demo'],
                                                              'status': 'yes'}
                                ls.append(a)
                                p += 0
                lss.append(p)

            print(ls)
            # print(len(ls))
            return render_template('a.html', data=ls)
    elif selectp == "title":
        res1 = db.select(
            " SELECT project.*,student.s_name,college.c_name,department.dept_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN department ON department.dept_id=student.dept_id where department.dept_id='"+str(session["cdepid"])+"' and  project.`title` LIKE '%" + search + "%' and (project.`p_year` between '" + yr1 + "' and '" + yr2 + "')")

        return render_template('a.html', data=res1)
    elif selectp == "department":
        res1 = db.select(
            " SELECT project.*,student.s_name,college.c_name,department.dept_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN department ON department.dept_id=student.dept_id where department.dept_id='"+str(session["cdepid"])+"' and  department.`dept_name` LIKE '%" + search + "%' and (project.`p_year` between '" + yr1 + "' and '" + yr2 + "')")

        return render_template('a.html', data=res1)
    elif selectp == "parea":
        res1 = db.select(
            " SELECT project.*,student.s_name,college.c_name,department.dept_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN department ON department.dept_id=student.dept_id where  department.dept_id='"+str(session["cdepid"])+"' and  project.`project_area` LIKE '%" + search + "%' and (project.`p_year` between '" + yr1 + "' and '" + yr2 + "') order by project.p_year desc")

        return render_template('a.html', data=res1)

@app.route('/searchpp1', methods=["post"])
def searchpp1():
    session["cdepid"]
    selectp = request.form['select']
    search = request.form['textbox10']
    yr1 = request.form['select2']
    yr2 = request.form['select3']
    db = Db()
    if selectp == "all":
        ss = search.split(" ")
        print(ss)
        res = db.select("SELECT project.*,student.s_name,college.c_name,project.p_year,department.dept_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN  department on student.dept_id=department.dept_id where department.dept_id='"+str(session["cdepid"])+"' and  project.`title` LIKE '%" + search + "%' and (project.`p_year` between '" + yr1 + "' and '" + yr2 + "')")
        pid = []
        ls = []
        for i2 in ss:
            if i2 != '':
                w = "SELECT distinct `key_words`.`pid`, project.*,student.s_name,college.c_name,department.dept_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN `key_words`  ON `project`.`project_id`=`key_words`.`pid` INNER JOIN department ON department.dept_id=student.dept_id WHERE department.dept_id='"+str(session["cdepid"])+"' and  (project.p_year between '" + yr1 + "' and '" + yr2 + "') AND (`key_words`.`key` LIKE '%" + i2 + "%') and `count`>1 ORDER BY `count`,p_year DESC"
                res = db.select(w)
                print(w)

                if len(res) > 0:

                    for i in res:
                        if i["project_id"] in pid:
                            print("---")
                            pass
                        else:

                            a = {'title': i['title'], 'abstract': i['abstract'], 'project_id': i['project_id'],
                                 's_name': i['s_name'], 'project_area': i['project_area'], 'dept_name': i['dept_name'],
                                 'c_name': i['c_name'], 'p_year': i['p_year'], 'report': i['report'],
                                 'journal': i['journal'], 'githublink': i['githublink'], 'demo': i['demo'],
                                 'status': 'yes'}
                            ls.append(a)

        return render_template('searchpp1.html', data=res)
    elif selectp == "title":
        res1 = db.select(
            " SELECT project.*,student.s_name,college.c_name,department.dept_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN department ON department.dept_id=student.dept_id where department.dept_id='"+str(session["cdepid"])+"' and  project.`title` LIKE '%" + search + "%' and (project.`p_year` between '" + yr1 + "' and '" + yr2 + "')")

        return render_template('searchpp1.html', data=res1)
    elif selectp == "department":
        res1 = db.select(
            " SELECT project.*,student.s_name,college.c_name,department.dept_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN department ON department.dept_id=student.dept_id where department.dept_id='"+str(session["cdepid"])+"' and  department.`dept_name` LIKE '%" + search + "%' and (project.`p_year` between '" + yr1 + "' and '" + yr2 + "')")

        return render_template('searchpp1.html', data=res1)
    elif selectp == "parea":
        res1 = db.select(
            " SELECT project.*,student.s_name,college.c_name,department.dept_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN department ON department.dept_id=student.dept_id where  department.dept_id='"+str(session["cdepid"])+"' and  project.`project_area` LIKE '%" + search + "%' and (project.`p_year` between '" + yr1 + "' and '" + yr2 + "') order by project.p_year desc")

        return render_template('searchpp1.html', data=res1)

@app.route('/select_col')
def select_col():
    db=Db()
    res = db.select("SELECT * FROM `college`")
    print(res)
    return render_template('/choosedepartment.html', col=res)

@app.route('/dept_ajax2',methods=["post"])
def dept_ajax2():
    colid=request.form['college_id']
    db=Db()
    q="SELECT * FROM `department` WHERE `college_id`='"+colid+"'"
    res=db.select(q)
    print(res)
    print(q)
    return jsonify(status="ok",data=res)

@app.route('/choosedepartment_post',methods=["post"])
def choosedepartment_post():
    dept = request.form['select2']
    session["cdepid"]=dept
    db = Db()
    res = db.select("SELECT project.*,student.s_name,department.dept_name,branch.branch_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN department ON student.dept_id=department.dept_id INNER JOIN `branch` ON student.branch_id=branch.branch_id WHERE department.dept_id='" + str(dept) + "'")
    return render_template('/searchp.html',data=res)

@app.route('/searchp1')
def searchp1():
    return render_template('searchp1.html')

@app.route('/searchp1_post',methods=["post"])
def searchp1_post():
    pid_list=[]
    ls=[]
    selectp = request.form['select']
    search = request.form['textbox10']
    yr1 = request.form['select2']
    yr2 = request.form['select3']
    db = Db()
    if selectp == "all":
        restt = search.split(" ")
        print(restt)
        lss = []
        ls = []
        pid = []
        qry = "SELECT project.project_id,project.abstract,journal,project.report,githublink,demo,project.project_area,dept_name,project.title,student.s_name,college.c_name,project.p_year,department.dept_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN `ngram` ON `ngram`.`pid`=`project`.`project_id` inner join `department` ON `department`.`dept_id`=`student`.`dept_id` WHERE (ngram.`word` = '" + search + "' or project.`title` = '" + search + "') and  p_year between '"+yr1+"' and '"+yr2+"'"
        print(qry)
        db = Db()
        res = db.select(qry)
        print(len(res), "------")
        if len(res) > 0:
            print(res)
            print(len(res))
            for i in res:
                if i["project_id"] in pid:
                    print("---", i["project_id"])
                    pass
                else:
                    print(pid, "--------------")
                    pid.append(i["project_id"])
                    a = {'title': i['title'], 'abstract': i['abstract'], 'project_id': i['project_id'],
                                                  's_name': i['s_name'], 'project_area': i['project_area'],
                                                  'dept_name': i['dept_name'],
                                                  'c_name': i['c_name'], 'p_year': i['p_year'], 'report': i['report'],
                                                  'journal': i['journal'], 'githublink': i['githublink'], 'demo': i['demo'],
                                                  'status': 'yes'}
                    ls.append(a)
            return render_template('a.html', data=ls)
        else:
            # ss = db.select(" SELECT project.project_id,project.title,student.s_name,college.c_name,project.p_year FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id where project.`title` LIKE '%" + search + "%' or project.`p_year`LIKE '%" + search + "%'")
            ss = db.select(
                " SELECT project.project_id,project.abstract,project.project_area,journal,project.report,githublink,demo,dept_name,project.title,student.s_name,college.c_name,project.p_year FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id inner join `department` ON `department`.`dept_id`=`student`.`dept_id` where   p_year between '"+yr1+"' and '"+yr2+"'")
            print(ss)
            pid = []
            ls = []
            for i in ss:
                p = 0
                for j in restt:
                    if j != '':
                        Q = "SELECT distinct `key_words`.`pid`,key_words.count, project.project_id,project.title,student.s_name,college.c_name,project.p_year FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN `key_words`  ON `project`.`project_id`=`key_words`.`pid` where `key_words`.`pid`='" + str(
                            i["project_id"]) + "' and key_words.`key` ='" + str(j) + "'"
                        pp = db.select(Q)
                        print(Q)
                        if len(pp) > 0:

                            if i["project_id"] in pid:
                                print("---", i["project_id"])
                                pass
                            else:
                                print(pid, "--------------")
                                pid.append(i["project_id"])
                                a = {'title': i['title'], 'abstract': i['abstract'], 'project_id': i['project_id'],
                                                              's_name': i['s_name'], 'project_area': i['project_area'],
                                                              'dept_name': i['dept_name'],
                                                              'c_name': i['c_name'], 'p_year': i['p_year'], 'report': i['report'],
                                                              'journal': i['journal'], 'githublink': i['githublink'], 'demo': i['demo'],
                                                              'status': 'yes'}
                                ls.append(a)
                                p += 0
                lss.append(p)

            print(ls)
            # print(len(ls))
            return render_template('a.html', data=ls)

    elif selectp == "title":
        res1 = db.select(
            " SELECT project.*,student.s_name,college.c_name,department.dept_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN department ON department.dept_id=student.dept_id where project.`title` LIKE '%" + search + "%' and (project.`p_year` between '" + yr1 + "' and '" + yr2 + "')")

        return render_template('a.html', data=res1)
    elif selectp == "department":
        res1 = db.select(
            " SELECT project.*,student.s_name,college.c_name,department.dept_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN department ON department.dept_id=student.dept_id where department.`dept_name` LIKE '%" + search + "%' and (project.`p_year` between '" + yr1 + "' and '" + yr2 + "')")

        return render_template('a.html', data=res1)
    elif selectp == "parea":
        res1 = db.select(" SELECT project.*,student.s_name,college.c_name,department.dept_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN department ON department.dept_id=student.dept_id where project.`project_area` LIKE '%" + search + "%' and (project.`p_year` between '" + yr1 + "' and '" + yr2 + "') order by project.p_year desc")

        return render_template('a.html', data=res1)


@app.route('/searchpro_post',methods=["post"])
def searchpro_post():
    selectp = request.form['select']
    search = request.form['textbox10']
    yr1 = request.form['select2']
    yr2 = request.form['select3']
    db = Db()
    if selectp == "all":
        ss = search.split(" ")
        print(ss)
        res = db.select(
            " SELECT project.*,student.s_name,college.c_name,project.p_year FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN  department on student.dept_id=department.dept_id where project.`title` LIKE '%" + search + "%' and (project.`p_year` between '" + yr1 + "' and '" + yr2 + "')")

        pid = []
        ls = []
        for i2 in ss:
            if i2 != '':
                w = "SELECT distinct `key_words`.`pid`, project.*,student.s_name,college.c_name,department.dept_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN `key_words`  ON `project`.`project_id`=`key_words`.`pid` INNER JOIN department ON department.dept_id=student.dept_id WHERE (project.p_year between '" + yr1 + "' and '" + yr2 + "') AND ((`project`.title LIKE '%" + i2 + "%') OR (`project`.project_area LIKE '%" + i2 + "%') OR (`key_words`.`key` LIKE '%" + i2 + "%')) ORDER BY `count`,p_year DESC"
                res = db.select(w)
                print(w)

                if len(res) > 0:

                    for i in res:
                        if i["project_id"] in pid:
                            print("---")
                            pass
                        else:

                            a = {'title': i['title'], 'abstract': i['abstract'], 'project_id': i['project_id'],
                                 's_name': i['s_name'], 'project_area': i['project_area'], 'dept_name': i['dept_name'],
                                 'c_name': i['c_name'], 'p_year': i['p_year'], 'report': i['report'],
                                 'journal': i['journal'], 'githublink': i['githublink'], 'demo': i['demo'],
                                 'status': 'yes'}
                            ls.append(a)

        return render_template('searchpro.html', data=res)
    elif selectp=="title":
        res1 = db.select(" SELECT project.*,student.s_name,college.c_name,department.dept_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN department ON department.dept_id=student.dept_id where project.`title` LIKE '%" + search + "%' and (project.`p_year` between '"+yr1+"' and '"+yr2+"')")

        return render_template('searchpro.html',data=res1)
    elif selectp=="department":
        res1 = db.select(" SELECT project.*,student.s_name,college.c_name,department.dept_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN department ON department.dept_id=student.dept_id where department.`dept_name` LIKE '%" + search + "%' and (project.`p_year` between '"+yr1+"' and '"+yr2+"')")

        return render_template('searchpr.html',data=res1)
    elif selectp=="parea":
        res1 = db.select(" SELECT project.*,student.s_name,college.c_name,department.dept_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN department ON department.dept_id=student.dept_id where project.`project_area` LIKE '%" + search + "%' and (project.`p_year` between '"+yr1+"' and '"+yr2+"') order by project.p_year desc")

        return render_template('searchpr.html',data=res1)

    @app.route('/viewmorepro/<project_id>')
    def viewmorepro(project_id):
        session["pid_id"] = project_id
        db = Db()
        res = db.select(
            "SELECT project.*,student.s_name,department.dept_name FROM project INNER JOIN student ON student.login_id=project.stud_id INNER JOIN department ON department.dept_id=student.dept_id AND project.project_id='" + str(
                project_id) + "'")
        print(res)
        return render_template('viewmorepro.html', data=res)

@app.route('/viewmoreproject/<project_id>')
def viewmoreproject(project_id):
    db = Db()
    res = db.select("SELECT project.*,student.s_name,department.dept_name FROM project INNER JOIN student ON student.login_id=project.stud_id INNER JOIN department ON department.dept_id=student.dept_id AND project.project_id='"+str(project_id)+"'")
    print(res)
    return render_template('viewmoreproject.html', data=res)

@app.route('/searchpr')
def searchpr():
    return render_template('searchpr.html')

@app.route('/searchpr_post',methods=["post"])
def searchpr_post():
    selectp=request.form['select']
    search=request.form['textbox10']
    yr1=request.form['select2']
    yr2 = request.form['select3']
    db=Db()
    if selectp=="all":
        ss = search.split(" ")
        print(ss)
        # res = db.select(" SELECT project.*,student.s_name,college.c_name,department.dept_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN  department on student.dept_id=department.dept_id where project.`title` LIKE '%" + search + "%' and (project.`p_year` between '"+yr1+"' and '"+yr2+"')")

        pid = []
        ls = []
        for i2 in ss:
            if i2 != '':
                w = "SELECT distinct `key_words`.`pid`, project.*,student.s_name,college.c_name,dept_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN  department on student.dept_id=department.dept_id INNER JOIN `key_words`  ON `project`.`project_id`=`key_words`.`pid` WHERE ( `project`.title LIKE '%" + i2 + "%') OR ( `project`.project_area LIKE '%" + i2 + "%') OR (`key_words`.`key` LIKE '%" + i2 + "%')  AND (project.p_year between '" + yr1 + "' and '" + yr2 + "')    group by project.project_id ORDER BY `count` DESC"
               # w = "SELECT distinct `key_words`.`pid`, project.*,student.s_name,college.c_name,department.dept_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN `key_words`  ON `project`.`project_id`=`key_words`.`pid` INNER JOIN department ON department.dept_id=student.dept_id WHERE (project.p_year between '" + yr1 + "' and '" + yr2 + "') AND ((`project`.title LIKE '%" + i2 + "%') OR (`project`.project_area LIKE '%" + i2 + "%') OR (`key_words`.`key` LIKE '%" + i2 + "%')) ORDER BY `count`,p_year DESC"
                res = db.select(w)
                print(w)

                if len(res) > 0:

                    for i in res:
                        if i["project_id"] in pid:
                            print("---")
                            pass
                        else:

                            a = {'title': i['title'], 'abstract': i['abstract'], 'project_id': i['project_id'],
                                 's_name': i['s_name'], 'project_area': i['project_area'], 'dept_name': i['dept_name'],
                                 'c_name': i['c_name'], 'p_year': i['p_year'], 'report': i['report'],
                                 'journal': i['journal'], 'githublink': i['githublink'], 'demo': i['demo'],
                                 'status': 'approved'}
                            ls.append(a)

        return render_template('b.html', data=ls)
    elif selectp=="title":
        res1 = db.select(" SELECT project.*,student.s_name,college.c_name,department.dept_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN department ON department.dept_id=student.dept_id where project.`title` LIKE '%" + search + "%' and (project.`p_year` between '"+yr1+"' and '"+yr2+"')")

        return render_template('b.html',data=res1)
    elif selectp=="department":

        res1 = db.select(" SELECT project.*,student.s_name,college.c_name,department.dept_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN department ON department.dept_id=student.dept_id where department.`dept_name` LIKE '%" + search + "%' and (project.`p_year` between '"+yr1+"' and '"+yr2+"')")

        return render_template('b.html',data=res1)
    elif selectp=="parea":
        res1 = db.select(" SELECT project.*,student.s_name,college.c_name,department.dept_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN department ON department.dept_id=student.dept_id where project.`project_area` LIKE '%" + search + "%' and (project.`p_year` between '"+yr1+"' and '"+yr2+"') order by project.p_year desc")

        return render_template('b.html',data=res1)

@app.route('/searchp',methods=["post"])
def searchp():
    selectp=request.form['select']
    search=request.form['textbox10']
    yr1=request.form['select2']
    yr2 = request.form['select3']
    db=Db()
    if selectp=="all":
        ss = search.split(" ")
        print(ss)
        # res = db.select(" SELECT project.*,student.s_name,college.c_name,project.p_year FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN  department on student.dept_id=department.dept_id where project.`title` LIKE '%" + search + "%' and (project.`p_year` between '"+yr1+"' and '"+yr2+"')")

        pid = []
        ls = []
        for i2 in ss:
            if i2 != '':
                w = "SELECT distinct `key_words`.`pid`, project.*,student.s_name,college.c_name,department.dept_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN `key_words`  ON `project`.`project_id`=`key_words`.`pid` INNER JOIN department ON department.dept_id=student.dept_id WHERE  ((`project`.title LIKE '%" + i2 + "%') OR (`project`.project_area LIKE '%" + i2 + "%') OR (`key_words`.`key` LIKE '%" + i2 + "%')) (project.p_year between '" + yr1 + "' and '" + yr2 + "') ORDER BY `count`,p_year DESC"
                res = db.select(w)
                print(w)

                if len(res) > 0:

                    for i in res:
                        if i["project_id"] in pid:
                            print("---")
                            pass
                        else:

                            a = {'title': i['title'], 'abstract': i['abstract'], 'project_id': i['project_id'],
                                 's_name': i['s_name'], 'project_area': i['project_area'], 'dept_name': i['dept_name'],
                                 'c_name': i['c_name'], 'p_year': i['p_year'], 'report': i['report'],
                                 'journal': i['journal'], 'githublink': i['githublink'], 'demo': i['demo'],
                                 'status': 'yes'}
                            ls.append(a)

        return render_template('searchp.html', data=res)
    elif selectp=="title":
        res1 = db.select(" SELECT project.*,student.s_name,college.c_name,project.p_year FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id where project.`title` LIKE '%" + search + "%' and (project.`p_year` between '"+yr1+"' and '"+yr2+"')")

        return render_template('searchp.html',data=res1)
    elif selectp=="title":
        res1 = db.select(" SELECT project.*,student.s_name,college.c_name,project.p_year FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id where project.`title` LIKE '%" + search + "%' and (project.`p_year` between '"+yr1+"' and '"+yr2+"')")

        return render_template('searchp.html',data=res1)


@app.route('/all_login')
def all_login():
    return render_template('all_login.html')

@app.route('/all_login_post', methods=['post'])
def all_login_post():
   db=Db()
   uname = request.form['username']
   password = request.form['password']
   q="select * from login WHERE username='"+uname+"'and password='"+password+"'"
   print(q)
   ss=db.selectOne(q)
   print(ss)

   if ss is not None :
       session["lid"]=ss["login_id"]
       # session["pid"]=ss["project_id"]
       if ss ['user_type']=='admin':
            return render_template('searchp1.html')
       elif ss ['user_type']=='college':
            db=Db()
            qry="SELECT * FROM college WHERE login_id='"+str(ss['login_id'])+"'"
            res1 = db.selectOne(qry)
            print(res1)
            session["name"]=res1["c_name"]
            session["photo"]=res1["image"]
            return render_template('searchp1.html')

       elif ss ['user_type']=='staff':
            qry2="SELECT * FROM `staff` WHERE `login_id`='"+str(ss["login_id"])+"'"
            res2=db.selectOne(qry2)
            print(res2)
            session["col_id"] = res2["college_id"]
            cid=res2["college_id"]
            qry = "SELECT * FROM college WHERE login_id='" + str(cid) + "'"
            res4 = db.selectOne(qry)
            print("-----------------")
            print(res4)
            session["clg"] = res4["c_name"]
            session["photo"] = res4["image"]
            return render_template('searchp1.html')

       elif ss ['user_type']=='student':
            qry2="SELECT * FROM `student` WHERE `login_id`='"+str(ss["login_id"])+"'"
            res2=db.selectOne(qry2)
            print(res2)
            session["col_id"] = res2["college_id"]
            cid=res2["college_id"]
            qry = "SELECT * FROM college WHERE login_id='" + str(cid) + "'"
            res4 = db.selectOne(qry)
            print("-----------------")
            print(res4)
            session["clg"] = res4["c_name"]
            session["photo"] = res4["image"]

            return render_template('searchp1.html')
       else:
            return '''<script>alert('Invalid Userid or password ');window.location='/all_login'</script>'''

   else:
        return "return '''<script>alert('Invalid Userid or password');window.location='/all_login'</script>'''"

@app.route('/allviewhome')
def allviewhome():
    db = Db()
    res = db.select( "SELECT project.*,student.s_name,department.dept_name FROM project INNER JOIN student ON student.login_id=project.stud_id INNER JOIN department ON department.dept_id=student.dept_id")
    print(res)
    return render_template('allviewhome.html',data=res)


@app.route('/allviewhome_post', methods=['post'])
def allviewhome_post():
    dept = request.form["dept"]
    btn = request.form["button10"]
    search1 = request.form['textbox10']
    yr1 = request.form['textbox1']
    yr2 = request.form['textbox2']
    db = Db()
    # ==========================
    ss = search1.split(" ")
    print(ss)

    pid = []
    ls = []
    for i2 in ss:
        if i2 != '':
            w = "SELECT distinct `key_words`.`pid`, project.*,student.s_name,college.c_name,department.dept_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN `key_words`  ON `project`.`project_id`=`key_words`.`pid` INNER JOIN department ON department.dept_id=student.dept_id WHERE  (project.p_year between '"+yr1+"' and '"+yr2+"') AND ((`project`.title LIKE '%" + i2 + "%') OR (`project`.project_area LIKE '%" + i2 + "%') OR  AND department.dept_name = '" + str(dept) + "'OR ((`key_words`.`key` LIKE '%" + i2 + "%')) ORDER BY `count` DESC)"
            res = db.select(w)
            print(w)

            if len(res) > 0:

                for i in res:
                    if i["project_id"] in pid:
                        print("---")
                        pass
                    else:

                        a = {'title': i['title'], 'abstract': i['abstract'], 'project_id': i['project_id'],
                             's_name': i['s_name'], 'project_area': i['project_area'], 'dept_name': i['dept_name'],
                             'c_name': i['c_name'], 'p_year': i['p_year'], 'report': i['report'], 'journal': i['journal'], 'githublink': i['githublink'], 'demo': i['demo'], 'status': 'yes'}
                        ls.append(a)

    return render_template('allviewhome.html', data=res)


@app.route('/index')
def index():
    return render_template('index.html')

@app.route('/login_post', methods=['post'])
def login_post():
   db=Db()
   uname = request.form['username']
   password = request.form['password']
   q="select * from login WHERE username='"+uname+"'and password='"+password+"'"
   print(q)
   ss=db.selectOne(q)

   print(ss)

   if ss is not None :
       session["lid"]=ss["login_id"]

       if ss ['user_type']=='admin':
            return render_template('adminhome.html')
       elif ss ['user_type']=='college':
            db=Db()
            qry="SELECT * FROM college WHERE login_id='"+str(ss['login_id'])+"'"
            res1 = db.selectOne(qry)
            print(res1)
            session["name"]=res1["c_name"]
            session["photo"]=res1["image"]
            return render_template('collegehome.html')

       elif ss ['user_type']=='staff':
            qry2="SELECT * FROM `staff` WHERE `login_id`='"+str(ss["login_id"])+"'"
            res2=db.selectOne(qry2)
            print(res2)
            session["staffname"] = res2["name"]
            session["col_id"] = res2["college_id"]
            cid=res2["college_id"]
            session["depid"] = res2["dept_id"]
            s="SELECT * FROM department where dept_id='"+str(session["depid"])+"'"
            res3 = db.selectOne(s)
            print("================")
            session["dname"] = res3["dept_name"]
            print("------*******************")
            qry = "SELECT * FROM college WHERE login_id='" + str(cid) + "'"
            res4 = db.selectOne(qry)
            print("-----------------")
            print(res4)
            session["clg"] = res4["c_name"]
            session["photo"] = res4["image"]

            return render_template('staffhome.html')

       else:
            return '''<script>alert('Invalid Userid or password ');window.location='/index'</script>'''

   else:
        return "return '''<script>alert('Invalid Userid or password');window.location='/index'</script>'''"

@app.route('/adminhome')
def adminhome():
    return render_template('/adminhome.html')
@app.route('/ahome')
def ahome():
    return render_template('admin/ahome.html')

@app.route('/college_registration')
def college_registration():
    return render_template('admin/college_registration.html')

@app.route("/spro")
def spro():
    return render_template("searchpro.html")
@app.route('/college_registration_post', methods=['post'])
def college_registration_post():
    db = Db()
    name = request.form['textfield']
    place = request.form['textfield2']
    post = request.form['textfield3']
    pin = request.form['textfield5']
    state = request.form['state']
    district = request.form['district']
    phone = request.form['textfield7']
    email = request.form['textfield8']
    website = request.form['textfield9']
    image = request.files['fileField']
    dt = time.strftime("%Y%m%d-%H%M%S")
    image.save(staticpath+"college_img\\"+dt+".jpg")
    path = "/static/college_img/"+dt+".jpg"
    db=Db()
    import random
    passw=str(random.randint(00000,99999))
    lid=db.insert("insert into login(username,password,user_type) VALUES('"+email+"','"+passw+"','college')")
    db.insert("insert into college (c_name,place,post,district,pin,c_state,phone,email,website,image,login_id) VALUES ('" + name + "','" + place + "','" + post + "','"+district+"','"+pin+"','"+state+"','"+phone+"','"+email+"','"+website+"','"+path+"','"+str(lid)+"')")

    # s = smtplib.SMTP(host='smtp.gmail.com', port=587)
    # s.starttls()
    # s.login("prolibtech001@gmail.com", "Prolibtech@123")
    # msg = MIMEMultipart()  # create a message.........."
    # message = "Messege from DNTL"
    # msg['From'] = "prolibtech001@gmail.com"
    # msg['To'] = email
    # msg['Subject'] = "Your Password for Drive Now Text Later"
    # body = "Your Account has been verified by our team. You Can login using your password - " + str(passw)
    # msg.attach(MIMEText(body, 'plain'))
    # s.send_message(msg)
    return '''<script>alert('College Registered succcessfully');window.location='/college_registration'</script>'''


@app.route('/view_colleges')
def view_colleges():
    db = Db()
    res = db.select("select * from `college`")
    print(res)
    return render_template('admin/view_colleges.html',data=res)



@app.route('/view_colleges_post', methods=['post'])
def view_colleges_post():
    search = request.form['search_college']
    db = Db()
    res = db.select("select * from `college` where c_name like '%"+search+"%'")
    print(res)
    return render_template('admin/view_colleges.html',data=res)


@app.route('/delete_college/<college_id>')
def delete_college(college_id):
    d = Db()
    qry = "delete from college where college_id='"+college_id+"'"
    res = d.delete(qry)
    return '''<script>alert('Deleted');window.location='/view_colleges'</script>'''

@app.route('/edit_college/<college_id>')
def edit_college(college_id):
    db=Db()
    qry = "SELECT * FROM `college` WHERE `college_id` ='"+str(college_id)+"'"
    res = db.selectOne(qry)
    return render_template('admin/edit_college.html',data=res)


@app.route('/edit_college_post',methods=['post'])
def edit_college_post():
    id = request.form['h1']
    c_name = request.form['textfield']
    place = request.form['textfield2']
    post = request.form['textfield3']
    district = request.form['district']
    pin = request.form['textfield5']
    state = request.form['state']
    phone = request.form['textfield7']
    email = request.form['textfield8']
    website = request.form['textfield9']
    image = request.files['fileField']
#     # dt = time.strftime("%Y%m%d-%H%M%S")
#     # image.save(staticpath + "college_img\\" + dt + ".jpg")
#     # path = "/static/college_img/" + dt + ".jpg"
    db=Db()
    if 'fileField' in request.files:
        image = request.files['fileField']
        if image.filename!="":
            dt = time.strftime("%Y%m%d-%H%M%S")
            image.save(staticpath + "college_img\\" + dt + ".jpg")
            path = "/static/college_img/" + dt + ".jpg"

            qry="UPDATE `college` SET c_name='"+c_name+"',place='"+place+"', post='"+post+"',district='"+district+"', c_state='"+state+"', pin='"+pin+"',phone='"+phone+"',email='"+email+"',website='"+website+"',image='"+path+"' WHERE college_id='"+str(id)+"'"
            res=db.update(qry)
            return '''<script>alert('Updated');window.location='/view_colleges'</script>'''
        else:
            qry = "UPDATE `college` SET c_name='" + c_name + "',place='" + place + "', post='" + post + "',district='" + district + "', c_state='" + state + "', pin='" + pin + "',phone='" + phone + "',email='" + email + "',website='" + website + "' WHERE college_id='"+str(id)+"'"
            res = db.update(qry)
            return '''<script>alert('Updated');window.location='/view_colleges'</script>'''
    else:
        qry = "UPDATE `college` SET c_name='" +c_name + "',place='" + place + "', post='" + post + "',district='" + district + "', c_state='" + state + "', pin='" + pin + "',phone='" + phone + "',email='" + email + "',website='" + website + "' WHERE college_id='"+str(id)+"'"
        res = db.update(qry)
        return '''<script>alert('Updated');window.location='/view_colleges'</script>'''
@app.route('/adddep')
def adddep():
    return render_template('admin/adddep.html')

@app.route('/adddep_post',methods=['post'])
def adddep_post():
    db=Db()
    dname = request.form['textfield']
    # image = request.files['fileField']
    # dt = time.strftime("%Y%m%d-%H%M%S")
    # image.save(staticpath + "d\\" + dt + ".jpg")
    # path = "/static/d/" + dt + ".jpg"
    db.insert("insert into dept(dpname) VALUES ('" + dname + "')")
    return '''<script>alert('Department added successfully');window.location='/adddep'</script>'''

@app.route('/viewdep')
def viewdep():
    db = Db()
    res = db.select("select * from `dept`")
    print(res)
    return render_template('admin/viewdep.html',data=res)
@app.route('/deldept/<dpno>')
def deldept(dpno):
    db = Db()
    qry = "delete from dept where dpno='" + dpno + "'"
    res = db.delete(qry)
    return '''<script>alert('Deleted');window.location='/viewdep'</script>'''
@app.route('/editdept/<dpno>')
def editdept(dpno):
    db = Db()
    qry = "SELECT * FROM `dept` WHERE `dpno` ='" + str(dpno) + "'"
    res = db.selectOne(qry)
    return render_template('admin/editdept.html', data=res)

@app.route('/editdept_post', methods=['post'])
def editdept_post():
    id = request.form['h1']
    dpname = request.form['textfield']
    db = Db()
    if 'fileField' in request.files:
        image = request.files['fileField']
        if image.filename!="":
            dt = time.strftime("%Y%m%d-%H%M%S")
            image.save(staticpath + "d\\" + dt + ".jpg")
            path = "/static/h/" + dt + ".jpg"
            qry="UPDATE `dept` SET dpname='"+dpname+"' and image='"+path+"'"
            res=db.update(qry)
            print(qry)
            return '''<script>alert('Updated');window.location='/viewdep'</script>'''
        else:

            qry = "UPDATE `dept` SET dpname='" + dpname + "'"
            res = db.update(qry)
            print(qry)
            return '''<script>alert('Updated');window.location='/viewdep'</script>'''
    else:
        qry = "UPDATE `dept` SET dpname='" + dpname + "'"
        res = db.update(qry)
        return '''<script>alert('Updated00');window.location='/viewdep'</script>'''
@app.route('/view_department/<id>')
def view_department(id):
    db = Db()
    res = db.select("select * from `department` where department.college_id='"+str(id)+"'")
    print(res)
    return render_template('admin/view_department.html',data=res)

@app.route('/view_staff/<dept_id>')
def view_staff(dept_id):
    db = Db()
    res = db.select("SELECT staff.* FROM staff INNER JOIN department ON staff.dept_id=department.dept_id AND department.dept_id='"+str(dept_id)+"'")
    print(res)
    return render_template('admin/view_staff.html', data=res)

@app.route('/viewview_project')
def viewview_project():
    db = Db()
    res = db.select("SELECT project.*,student.s_name,department.dept_name FROM project INNER JOIN student ON student.login_id=project.stud_id INNER JOIN department ON department.dept_id=student.dept_id")
    print(res)
    return render_template('admin/viewview_project.html', data=res)

@app.route('/report_keyword_extraction/<pid>')
def report_keyword_extraction(pid):
    q="SELECT `report` FROM `project` WHERE `project_id`='"+pid+"'"
    d=Db()
    res=d.selectOne(q)
    if res is not None:
        import PyPDF2, yake
        pdfFileObj = open('C:\\Users\\User\\PycharmProjects\\project_management12\\static\\report\\20220509-021543.pdf',
                          'rb')
        # creating a pdf reader object
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
        # printing number of pages in pdf file
        print(pdfReader.numPages)
        # creating a page object
        gg = ""
        for i in range(0, pdfReader.numPages):
            print(i)
            pageObj = pdfReader.getPage(i)
            gg = gg + " " + pageObj.extractText()
            # extracting text from page
            text = pageObj.extractText()
            print(text)
            language = "en"
            max_ngram_size = 3
            deduplication_threshold = 0.9
            numOfKeywords = 10
            custom_kw_extractor = yake.KeywordExtractor(lan=language, n=max_ngram_size, top=numOfKeywords,
                                                        features=None)
            keywords = custom_kw_extractor.extract_keywords(gg)

            pdfFileObj.close()
            a = gg.replace(",", "").replace(".", "").replace("0", "").replace("1", "").replace("2", "").replace("3",
                                                                                                                "").replace(
                "4", "").replace("5", "").replace("6", "").replace("7", "").replace("8", "").replace("0", "").replace(
                "!",
                "").replace(
                "@", "").replace('-', "")

            kw_extractor = yake.KeywordExtractor()
            language = "en"
            max_ngram_size = 3
            deduplication_threshold = 0.9
            numOfKeywords = 200
            custom_kw_extractor = yake.KeywordExtractor(lan=language, n=max_ngram_size, top=numOfKeywords,
                                                        features=None)
            keywords = custom_kw_extractor.extract_keywords(a)
            for kw in keywords:
                print(kw[0])
                print(kw[1])

    return render_template('admin/words.html')

@app.route('/report_word_extraction/<pid>')
def report_word_extraction(pid):
    q="SELECT `report` FROM `project` WHERE `project_id`='"+pid+"'"
    d=Db()
    res=d.selectOne(q)
    if res is not None:
        import PyPDF2

        # creating a pdf file object
        pdfFileObj = open(
            'C:\\Users\\User\\PycharmProjects\\project_management12\\static\\report\\20220429-104419.pdf', 'rb')

        # creating a pdf reader object
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj)

        # printing number of pages in pdf file
        print(pdfReader.numPages)

        # creating a page object
        gg=""
        for i in range(0, pdfReader.numPages):
            print(i)
            pageObj = pdfReader.getPage(i)
            gg=gg+" "+pageObj.extractText()

            # extracting text from page
            print(pageObj.extractText())

        # closing the pdf file object
        pdfFileObj.close()
        gg.replace(",", "").replace(".", "").replace("0", "").replace("1", "").replace("2", "").replace("3",
                                                                                                              "").replace(
            "4", "").replace("5", "").replace("6", "").replace("7", "").replace("8", "").replace("0", "").replace("!",
                                                                                                                  "").replace(
            "@", "")
        ss = word_tokenize(gg)
        print(ss)
        filtered_sentence = []

    for w in ss:
        if w not in stop_words:
            filtered_sentence.append(w)

    print(ss)
    print(filtered_sentence)
    # aa = list(session["wordslist"])
    # print(aa)
    aa=[]
    for i in filtered_sentence:
        aa.append(i)
    session["wordslist"] = aa
    print(aa)
    return render_template('admin/words.html',words=aa)

@app.route('/viewview_project_post', methods=['post'])
def viewview_project_post():
    btn=request.form["button"]
    if(btn=="Submit"):
        year = request.form['select3']
        db = Db()
        # qry="select student.*,department.dept_name,branch.branch_name FROM `department` INNER JOIN `student` ON `student`.`dept_id`=`department`.`dept_id` INNER JOIN `branch` ON `student`.`branch_id`= `branch`.`branch_id` WHERE student.`branch_id`='"+str(session['branch_id'])+"' AND `student`.`s_year` = '"+s_year+"' "
        res = db.select("SELECT project.*,student.s_name,department.dept_name,branch.branch_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN department ON student.dept_id=department.dept_id INNER JOIN `branch` ON student.branch_id=branch.branch_id WHERE `project`.`p_year` = '"+year+"'")

    elif btn=="view":
        year = request.form['select3']
        search = request.form['textbox10']
        db = Db()
        # qry = "select student.*,department.dept_name,branch.branch_name FROM `department` INNER JOIN `student` ON `student`.`dept_id`=`department`.`dept_id` INNER JOIN `branch` ON `student`.`branch_id`= `branch`.`branch_id` WHERE student.`branch_id`='" + str(session['branch_id'])+"' and `student`.`s_name` LIKE '%" + search + "%'"
        res = db.select("SELECT project.*,student.s_name,department.dept_name,branch.branch_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN department ON student.dept_id=department.dept_id INNER JOIN `branch` ON student.branch_id=branch.branch_id WHERE `project`.`title` LIKE '%" + search + "%' and `project`.`p_year` LIKE '%" + year + "%'")
    print("------------------------------")
    return render_template('admin/viewview_project.html',data=res)


# # ----------------------------COLLEGE----------------------------------------------------
@app.route('/collegehome')
def college():
    return render_template('/collegehome.html')
@app.route('/college_profile')
def college_profile():
    db = Db()
    qry = "SELECT * FROM `college` WHERE `login_id` ='" +str(session['lid'])+ "'"
    res = db.selectOne(qry)
    return render_template('college/college_profile.html', data=res)

@app.route('/college_changepwd')
def college_changepwd():
    return render_template('college/college_changepwd.html')

@app.route('/college_changepwd_post',methods=["post"])
def college_changepwd_post():
    db=Db()
    newpwd = request.form['textfield2']
    confnewpwd = request.form['textfield3']
    if newpwd==confnewpwd:
        qry = "UPDATE `login` SET password='" + newpwd + "' where login_id ='" +str(session["lid"])+ "'"
        res = db.update(qry)
        return '''<script>alert('updated');window.location='/index'</script>'''
    else:
        return '''<script>alert('Password mismatch');window.location='/college_changepwd'</script>'''
    # else:
    #     return '''<script>alert('Current password must be valid');window.location='/college_changepwd'</script>'''



@app.route('/staff_changepwd')
def staff_changepwd():
    return render_template('staff/staff_changepwd.html')

@app.route('/staff_changepwd_post',methods=["post"])
def staff_changepwd_post():
    db = Db()
    newpwd = request.form['textfield2']
    confnewpwd = request.form['textfield3']

    if newpwd==confnewpwd:
        qry = "UPDATE `login` SET password='" + newpwd + "' where login_id ='" +str(session["lid"])+ "'"
        res = db.update(qry)
        return '''<script>alert('updated');window.location='/index'</script>'''
    else:
        return '''<script>alert('Password mismatch');window.location='/staff_changepwd'</script>'''

@app.route('/view_department1')
def view_department1():
    db = Db()
    res = db.select("SELECT * FROM `department` WHERE `college_id`='"+str(session["lid"])+"'")
    print(res)
    return render_template('college/view_department1.html',data=res)

@app.route('/view_studdepartment1')
def view_studdepartment1():
    db = Db()
    res = db.select("SELECT * FROM `department` WHERE `college_id`='"+str(session["lid"])+"'")
    print(res)
    return render_template('college/view_studdepartment1.html',data=res)

@app.route('/view_studbranch/<dept_id>')
def view_studbranch(dept_id):
    db = Db()
    session['depid']=dept_id
    res = db.select("SELECT * FROM `branch`,`department` WHERE `branch`.`dept_id` = `department`.`dept_id` and `branch`.`dept_id` = '"+str(dept_id)+"'")
    return render_template('college/view_studbranch.html',data=res)

@app.route('/delete_department/<dept_id>')
def delete_department(dept_id):
    db = Db()
    qry = "delete from department where dept_id='" + dept_id + "'"
    res = db.delete(qry)
    return '''<script>alert('Deleted');window.location='/view_department1'</script>'''


@app.route('/edit_department/<dept_id>')
def edit_department(dept_id):
    db = Db()
    qry = "SELECT * FROM `department` WHERE `dept_id` ='" + str(dept_id) + "'"
    res = db.selectOne(qry)
    return render_template('college/edit_department.html', data=res)


@app.route('/edit_department_post', methods=['post'])
def edit_department_post():
    id = request.form['h1']
    dept_name = request.form['textfield']
    db = Db()
    qry = "UPDATE `department` SET dept_name='" + dept_name + "' where dept_id ='" + str(id) + "'"
    res = db.update(qry)
    return '''<script>alert('updated');window.location='/view_department1'</script>'''

@app.route('/view_branch/<dept_id>')
def view_branch(dept_id):
    db = Db()
    session['depid']=dept_id
    res = db.select("SELECT * FROM `branch`,`department` WHERE `branch`.`dept_id` = `department`.`dept_id` and `branch`.`dept_id` = '"+str(dept_id)+"'")
    return render_template('college/view_branch.html',data=res)
@app.route('/edit_branch/<branch_id>')
def edit_branch(branch_id):
    db = Db()
    qry = "SELECT * FROM `branch` WHERE `branch_id` ='" + str(branch_id) + "'"
    res = db.selectOne(qry)
    return render_template('college/edit_branch.html', data=res)

@app.route('/edit_branch_post', methods=['post'])
def edit_branch_post():
    id = request.form['h1']
    branch_name = request.form['textfield']
    db = Db()
    qry = "UPDATE `branch` SET branch_name='" + branch_name + "' where branch_id ='" + str(id) + "'"
    res = db.update(qry)
    return '''<script>alert('Branch updated successfully');window.location='/view_branch'</script>'''
@app.route('/delete_branch/<branch_id>')
def delete_branch(branch_id):
    db = Db()
    qry = "delete from branch where branch_id='" + branch_id + "'"
    res = db.delete(qry)
    return '''<script>alert('Deleted');window.location='/view_department1'</script>'''

@app.route('/add_branch')
def add_branch():
    return render_template('college/add_branch.html')

@app.route('/add_branch_post',methods=['post'])
def add_branch_post():
    db = Db()
    depid=session['depid']
    branch_name = request.form['textfield']
    db.insert("insert into branch (dept_id,college_id,branch_name) VALUES ('"+depid+"','"+str(session['lid'])+"','" + branch_name +"')")
    return '''<script>alert('branch name added successfully');window.location='/add_branch'</script>'''


@app.route('/add_department')
def add_department():
    db = Db()
    return render_template('college/add_department.html')

@app.route('/add_department_post',methods=['post'])
def add_department_post():
    db = Db()
    dpname = request.form['dept']
    # qry ="SELECT dept.dpname FROM department inner join dept on  department.dpno=dept.dpno  and dept.dpno='"+str(dpid)+"'"
    # dpname=db.selectOne(qry)
    # print(dpname)
    db.insert("insert into department (college_id,dept_name) VALUES ('"+str(session['lid'])+"','" + str(dpname) +"')")
    return '''<script>alert('department added successfully');window.location='/add_department'</script>'''

@app.route('/staff_registration')
def staff_registration():
    db=Db()
    res=db.select("SELECT * FROM `department` WHERE college_id = '"+str(session["lid"])+"'")
    print(res)

    return render_template('college/staff_registration.html', dept=res)


@app.route('/staff_registration_post',methods=['post'])
def staff_registration_post():
    db = Db()
    name = request.form['textfield']
    email = request.form['textfield2']
    phone = request.form['textfield3']
    department = request.form['select1']
    image = request.files['fileField']
    dt = time.strftime("%Y%m%d-%H%M%S")
    image.save(staticpath + "staff_img\\" + dt + ".jpg")
    path2 = "/static/staff_img/" + dt + ".jpg"

    import random
    passw = str(random.randint(00000, 99999))
    lid=db.insert("insert into login(username,password,user_type) VALUES('" + email + "','" + passw + "','staff')")
    db.insert("insert into staff (college_id,`name`,email,phone,dept_id,image,login_id) VALUES ('"+str(session['lid'])+"','" + name + "','" + email + "','" + phone + "','" + department + "','" + path2 + "','"+str(lid)+"')")
    # s = smtplib.SMTP(host='smtp.gmail.com', port=587)
    # s.starttls()
    # s.login("prolibtech001@gmail.com", "Prolibtech@123")
    # msg = MIMEMultipart()  # create a message.........."
    # message = "Messege from DNTL"
    # msg['From'] = "prolibtech001@gmail.com"
    # msg['To'] = email
    # msg['Subject'] = "Your Password for Drive Now Text Later"
    # body = "Your Account has been verified by our team. You Can login using your password - " + str(passw)
    # msg.attach(MIMEText(body, 'plain'))
    # s.send_message(msg)
    return '''<script>alert('staff added successfully');window.location='/staff_registration'</script>'''

@app.route('/view_staff1')
def view_staff1():
    db = Db()
    res = db.select("SELECT * FROM `staff` WHERE `college_id` ='"+str(session['lid'])+"'")
    print(res)
    dept=db.select("SELECT * FROM `department` WHERE college_id = '" + str(session["lid"]) + "'")
    return render_template('college/view_staff.html',data=res,dept=dept)


@app.route('/view_staff_post', methods=['post'])
def view_staff_post1():
    search = request.form['select1']
    db = Db()
    res = db.select("select * from  `staff` WHERE college_id = '" + str(session["lid"]) + "' AND dept_id='"+search+"'")
    print(res)

    dept = db.select("SELECT * FROM `department` WHERE college_id = '" + str(session["lid"]) + "'")
    return render_template('college/view_staff.html', data=res,dept=dept)

@app.route('/edit_staff/<staff_id>')
def edit_staff(staff_id):
    db = Db()
    qry = "SELECT * FROM `staff` WHERE `staff_id` ='" + str(staff_id) + "'"
    print(qry)
    res = db.selectOne(qry)
    r = db.select("SELECT * FROM `department` WHERE college_id = '" + str(session["lid"]) + "'")
    print(r)
    return render_template('college/edit_staff.html', data=res,dept=r)



@app.route('/edit_staff_post', methods=['post'])
def edit_staff_post():
    id = request.form['h1']
    name = request.form['textfield']
    email = request.form['textfield2']
    phone = request.form['textfield3']
    department = request.form['select1']
    db = Db()
    if 'fileField' in request.files:
        image = request.files['fileField']
        if image.filename!="":
            dt = time.strftime("%Y%m%d-%H%M%S")
            image.save(staticpath + "staff_img\\" + dt + ".jpg")
            path = "/static/staff_img/" + dt + ".jpg"
            qry="UPDATE `staff` SET name='"+name+"',phone='"+phone+"',email='"+email+"',dept_id='"+department+"',image='"+path+"' WHERE staff_id='"+str(id)+"'"
            res=db.update(qry)
            return '''<script>alert('Updated');window.location='/view_staff1'</script>'''
        else:
            qry = "UPDATE `staff` SET name='" + name + "',phone='" + phone + "',email='" + email + "',dept_id='" + department + "' WHERE staff_id='" + str(id) + "'"
            res = db.update(qry)
            return '''<script>alert('Updated');window.location='/view_staff1'</script>'''
    else:
        qry = "UPDATE `staff` SET name='" + name + "',phone='" + phone + "',email='" + email + "',dept_id='" + department + "' WHERE staff_id='" + str(id) + "'"
        res = db.update(qry)
        return '''<script>alert('Updated');window.location='/view_staff1'</script>'''

@app.route('/delete_staff/<staff_id>')
def delete_staff(staff_id):
    db = Db()
    qry = "delete from staff where staff_id='" + staff_id + "'"
    res = db.delete(qry)
    return '''<script>alert('Deleted');window.location='/view_staff1'</script>'''

@app.route('/student_registration')
def student_registration():
    db=Db()
    res=db.select("SELECT * FROM `department` WHERE college_id = '"+str(session["lid"])+"'")
    print(res)
    re = db.select("SELECT * FROM `branch` WHERE college_id = '" + str(session["lid"]) + "'")
    print(re)
    return render_template('college/student_registration.html', dept=res,branch=re)


@app.route('/student_registration_post',methods=['post'])
def student_registration_post():
    db = Db()
    name = request.form['textfield']
    email = request.form['textfield2']
    phone = request.form['textfield3']
    department = request.form['select1']
    branch = request.form['select2']
    year=request.form['select3']
    image = request.files['fileField']
    dt = time.strftime("%Y%m%d-%H%M%S")
    image.save(staticpath + "student\\" + dt + ".jpg")
    path = "/static/student/" + dt + ".jpg"
    import random
    passw = str(random.randint(0000, 9999))

    lid=db.insert("insert into login(username,password,user_type) VALUES('" + email + "','" + passw + "','student')")
    db.insert("insert into student (college_id,`s_name`,email,phone,dept_id,branch_id,`year`,image,login_id) VALUES ('"+str(session['lid'])+"','" + name + "','" + email + "','" + phone + "','" + department + "','" + branch + "','" + year + "','" + path + "','"+str(lid)+"')")
    # s = smtplib.SMTP(host='smtp.gmail.com', port=587)
    # s.starttls()
    # s.login("prolibtech001@gmail.com", "Prolibtech@123")
    # msg = MIMEMultipart()  # create a message.........."
    # message = "Messege from DNTL"
    # msg['From'] = "prolibtech001@gmail.com"
    # msg['To'] = email
    # msg['Subject'] = "Your Password for Drive Now Text Later"
    # body = "Your Account has been verified by our team. You Can login using your password - " + str(passw)
    # msg.attach(MIMEText(body, 'plain'))
    # s.send_message(msg)
    return '''<script>alert('student registered successfully');window.location='/student_registration'</script>'''


@app.route('/select_department')
def select_department():
    db=Db()
    res = db.select("SELECT * FROM `department` WHERE `college_id` ='" + str(session['lid']) + "'")
    print(res)
    return render_template('college/view_student.html', dept=res)
@app.route('/branch_ajax',methods=["post"])
def brach_ajax():
    depid=request.form['depid']
    db=Db()
    q="SELECT * FROM `branch` WHERE `dept_id`='"+depid+"'"
    res=db.select(q)
    print(res)
    print(q)

    return jsonify(status="ok",data=res)

@app.route('/select_branch/<dept_id>',methods=['post'])
def select_branch(dept_id):
    db=Db()
    session['depid'] = dept_id
    re = db.select("SELECT * FROM `branch`,`department` WHERE `branch`.`dept_id` = `department`.`dept_id` and `branch`.`dept_id` = '" + str(dept_id) + "'")
    print(re)
    return render_template('college/view_student.html', branch=re)

@app.route('/view_sstudent/<branch_id>')
def view_student(branch_id):
    db = Db()
    session['branch_id'] = branch_id
    # res = db.select("SELECT `student`.*,`department`.`dept_name`,`branch`.`branch_name` FROM `department` INNER JOIN `student` ON `student`.`dept_id`=`department`.`dept_id` INNER JOIN `branch` ON `student`.`branch_id`= `branch`.`branch_id` WHERE `student`.`college_id` ='"+str(session['lid'])+"' and `student`.`branch_id`='"+str(session['branch_id'])+"'")
    res=db.select("select student.*,department.dept_name,branch.branch_name FROM `department` INNER JOIN `student` ON `student`.`dept_id`=`department`.`dept_id` INNER JOIN `branch` ON `student`.`branch_id`= `branch`.`branch_id` WHERE student.`branch_id`='"+str(session['branch_id'])+"'" )
    print(res)
    return render_template('college/view_sstudent.html',data=res)

@app.route('/view_sstudent_post', methods=['post'])
def view_sstudent_post():
    btn=request.form["button"]
    if(btn=="Submit"):
        s_year = request.form['select3']

        db = Db()
        qry="select student.*,department.dept_name,branch.branch_name FROM `department` INNER JOIN `student` ON `student`.`dept_id`=`department`.`dept_id` INNER JOIN `branch` ON `student`.`branch_id`= `branch`.`branch_id` WHERE student.`branch_id`='"+str(session['branch_id'])+"' AND `student`.`year` = '"+s_year+"' "
        # qry="SELECT `student`.*,`department`.`dept_name`,`branch`.`branch_name` FROM `student` INNER JOIN branch ON `student`.`branch_id`=`branch`.`branch_id` INNER JOIN `department` ON `branch`.`dept_id`=`department`.`dept_id` WHERE `student`.`college_id` ='"+str(session['lid'])+"' AND `student`.`year` = '"+year+"' "
        res = db.select(qry)

    elif btn=="view":
        search = request.form['textbox10']
        db = Db()
        qry = "select student.*,department.dept_name,branch.branch_name FROM `department` INNER JOIN `student` ON `student`.`dept_id`=`department`.`dept_id` INNER JOIN `branch` ON `student`.`branch_id`= `branch`.`branch_id` WHERE student.`branch_id`='" + str(session['branch_id'])+"' and `student`.`s_name` LIKE '%" + search + "%'"

        # qry = "SELECT `student`.*,`department`.`dept_name`,`branch`.`branch_name` FROM `department` INNER JOIN `student` ON `student`.`dept_id`=`department`.`dept_id` INNER JOIN `branch` ON `student`.`branch_id`= `branch`.`branch_id` WHERE `student`.`s_name` LIKE '%" + search + "%' or `department`.`dept_name` LIKE+'%" + search +"%' or `student`.`year` LIKE+'%" + search +"%'"
        res = db.select(qry)

    print("------------------------------")
    print(qry)
    return render_template('college/view_sstudent.html',data=res)


@app.route('/edit_student/<stud_id>')
def edit_student(stud_id):
    db = Db()
    qry = "SELECT * FROM `student` WHERE `stud_id` ='" + str(stud_id) + "'"
    print(qry)
    res = db.selectOne(qry)
    dept = db.select("SELECT * FROM `department` WHERE college_id = '" + str(session["lid"]) + "'")
    print(dept)
    br = db.select("SELECT * FROM `branch` WHERE college_id = '" + str(session["lid"]) + "'")
    print(br)
    return render_template('college/edit_student.html',data=res,dept=dept,branch=br)


@app.route('/edit_student_post', methods=['post'])
def edit_student_post():
    id = request.form['h1']
    name = request.form['textfield']
    email = request.form['textfield2']
    phone = request.form['textfield3']
    department = request.form['select1']
    year=request.form["textfield4"]
    branch=request.form['select2']
    db = Db()
    if 'fileField' in request.files:
        image = request.files['fileField']
        if image.filename!="":
            dt = time.strftime("%Y%m%d-%H%M%S")
            image.save(staticpath + "student_img\\" + dt + ".jpg")
            path = "/static/student_img/" + dt + ".jpg"
            qry="UPDATE `student` SET s_name='"+name+"',phone='"+phone+"',email='"+email+"',dept_id='"+department+"',branch_id='"+branch+"',image='"+path+"',year='"+year+"' WHERE stud_id='"+str(id)+"'"
            res=db.update(qry)
            print(qry)
            return '''<script>alert('Updated');window.location='/view_student'</script>'''
        else:
            qry = "UPDATE `student` SET s_name='" + name + "',phone='" + phone + "',email='" + email + "',dept_id='" + department + "',branch_id='"+branch+"',year='"+year+"' WHERE stud_id='" + str(id) + "'"
            res = db.update(qry)
            return '''<script>alert('Updated');window.location='/view_student'</script>'''
    else:
        qry = "UPDATE `student` SET s_name='" + name + "',phone='" + phone + "',email='" + email + "',dept_id='" + department + "',branch_id='"+branch+"',year='"+year+"' WHERE stud_id='" + str(id) + "'"
        res = db.update(qry)
        return '''<script>alert('Updated');window.location='/view_student'</script>'''

@app.route('/delete_student/<stud_id>')
def delete_student(stud_id):
    db = Db()
    qry = "delete from student where stud_id='" + stud_id + "'"
    res = db.delete(qry)
    return '''<script>alert('Deleted');window.location='/view_student'</script>'''


@app.route('/view_prdepartment1')
def view_prdepartment1():
    db = Db()
    res = db.select("SELECT * FROM `department` WHERE `college_id`='"+str(session["lid"])+"'")
    print(res)
    return render_template('college/view_prdepartment1.html',data=res)
@app.route('/view_prdepartment2')
def view_prdepartment2():
    db = Db()
    res = db.select("SELECT department.* from department inner join college on department.college_id=college.login_id where department.college_id = '" +str(session['col_id'])+ "'")
    print(res)
    return render_template('student/view_prdepartment2.html',data=res)

@app.route('/view_prbranch/<dept_id>')
def view_prbranch(dept_id):
    db = Db()
    session['depid']=dept_id
    res = db.select("SELECT * FROM `branch`,`department` WHERE `branch`.`dept_id` = `department`.`dept_id` and `branch`.`dept_id` = '"+str(dept_id)+"'")
    return render_template('college/view_prbranch.html',data=res)

@app.route('/view_prbranch2/<dept_id>')
def view_prbranch2(dept_id):
    db = Db()
    session['depid']=dept_id
    res = db.select("SELECT * FROM `branch`,`department` WHERE `branch`.`dept_id` = `department`.`dept_id` and `branch`.`dept_id` = '"+str(dept_id)+"'")
    return render_template('student/view_prbranch2.html',data=res)

@app.route('/view_pprojects/<branch_id>')
def view_pprojects(branch_id):
    db=Db()
    session['branch_id'] = branch_id
    # qry="SELECT project.*,student.s_name,department.dept_name,branch.branch_name FROM project INNER JOIN student ON student.stud_id=project.stud_id INNER JOIN department ON student.dept_id=department.dept_id INNER JOIN `branch` ON student.branch_id=branch.branch_id WHERE student.branch_id='" + str(session['branch_id']) + "'"
    res=db.select("SELECT project.*,student.s_name,department.dept_name,branch.branch_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN department ON student.dept_id=department.dept_id INNER JOIN `branch` ON student.branch_id=branch.branch_id WHERE student.branch_id='" + str(session['branch_id']) + "'")

    return render_template('college/view_pprojects.html', data=res)

@app.route('/view_pprojects_post', methods=['post'])
def view_pprojects_post():
    btn=request.form["button"]
    if(btn=="Submit"):
        year = request.form['select3']
        db = Db()
        # qry="select student.*,department.dept_name,branch.branch_name FROM `department` INNER JOIN `student` ON `student`.`dept_id`=`department`.`dept_id` INNER JOIN `branch` ON `student`.`branch_id`= `branch`.`branch_id` WHERE student.`branch_id`='"+str(session['branch_id'])+"' AND `student`.`s_year` = '"+s_year+"' "
        res = db.select("SELECT project.*,student.s_name,department.dept_name,branch.branch_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN department ON student.dept_id=department.dept_id INNER JOIN `branch` ON student.branch_id=branch.branch_id WHERE student.branch_id='" + str(session['branch_id']) + "' AND `project`.`p_year` = '"+year+"'")

    elif btn=="view":
        year = request.form['select3']
        search = request.form['textbox10']
        db = Db()
        # qry = "select student.*,department.dept_name,branch.branch_name FROM `department` INNER JOIN `student` ON `student`.`dept_id`=`department`.`dept_id` INNER JOIN `branch` ON `student`.`branch_id`= `branch`.`branch_id` WHERE student.`branch_id`='" + str(session['branch_id'])+"' and `student`.`s_name` LIKE '%" + search + "%'"
        res = db.select("SELECT project.*,student.s_name,department.dept_name,branch.branch_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN department ON student.dept_id=department.dept_id INNER JOIN `branch` ON student.branch_id=branch.branch_id WHERE student.branch_id='" + str(session['branch_id']) + "' AND `project`.`p_year` = '"+year+"' AND `project`.`title` LIKE '%" + search + "%'")

    print("------------------------------")
    return render_template('college/view_pprojects.html',data=res)

@app.route('/view_pprojects2/<branch_id>')
def view_pprojects2(branch_id):
    db=Db()
    session['branch_id'] = branch_id
    # qry="SELECT project.*,student.s_name,department.dept_name,branch.branch_name FROM project INNER JOIN student ON student.stud_id=project.stud_id INNER JOIN department ON student.dept_id=department.dept_id INNER JOIN `branch` ON student.branch_id=branch.branch_id WHERE student.branch_id='" + str(session['branch_id']) + "'"
    res=db.select("SELECT project.*,student.s_name,department.dept_name,branch.branch_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN department ON student.dept_id=department.dept_id INNER JOIN `branch` ON student.branch_id=branch.branch_id WHERE student.branch_id='" + str(session['branch_id']) + "'")

    return render_template('student/view_pprojects2.html', data=res)

@app.route('/view_pprojects2_post', methods=['post'])
def view_pprojects2_post():
    btn=request.form["button"]
    if(btn=="Submit"):
        year = request.form['select3']
        db = Db()
        # qry="select student.*,department.dept_name,branch.branch_name FROM `department` INNER JOIN `student` ON `student`.`dept_id`=`department`.`dept_id` INNER JOIN `branch` ON `student`.`branch_id`= `branch`.`branch_id` WHERE student.`branch_id`='"+str(session['branch_id'])+"' AND `student`.`s_year` = '"+s_year+"' "
        res = db.select("SELECT project.*,student.s_name,department.dept_name,branch.branch_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN department ON student.dept_id=department.dept_id INNER JOIN `branch` ON student.branch_id=branch.branch_id WHERE student.branch_id='" + str(session['branch_id']) + "' AND `project`.`p_year` = '"+year+"'")

    elif btn=="view":
        year = request.form['select3']
        search = request.form['textbox10']
        db = Db()
        # qry = "select student.*,department.dept_name,branch.branch_name FROM `department` INNER JOIN `student` ON `student`.`dept_id`=`department`.`dept_id` INNER JOIN `branch` ON `student`.`branch_id`= `branch`.`branch_id` WHERE student.`branch_id`='" + str(session['branch_id'])+"' and `student`.`s_name` LIKE '%" + search + "%'"
        res = db.select("SELECT project.*,student.s_name,department.dept_name,branch.branch_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN department ON student.dept_id=department.dept_id INNER JOIN `branch` ON student.branch_id=branch.branch_id WHERE student.branch_id='" + str(session['branch_id']) + "' AND `project`.`p_year` = '"+year+"' AND `project`.`title` LIKE '%" + search + "%'")

    print("------------------------------")
    return render_template('student/view_pprojects2.html',data=res)

@app.route('/staffhome')
def staffhome():
    return render_template('staffhome.html')

@app.route('/staff_profile')
def staff_profile():
    db = Db()
    # qry="SELECT staff.*,department.dept_name FROM department INNER JOIN staff on staff.dept_id=department.dept_id WHERE staff.'login_id'= ='" +str(session['lid'])+ "'"
    qry = "SELECT * FROM `staff`,department WHERE `login_id` ='" +str(session['lid'])+ "'"
    res = db.selectOne(qry)
    return render_template('staff/staff_profile.html', data=res)

@app.route('/staffedit_profile/<staff_id>')
def staffedit_profile(staff_id):
    db = Db()
    qry = "SELECT * FROM `staff` WHERE `staff_id` ='" + str(staff_id) + "'"
    print(qry)
    res = db.selectOne(qry)
    return render_template('staff/staffedit_profile.html', data=res)



@app.route('/staffedit_profile_post', methods=['post'])
def staffedit_profile_post():
    id = request.form['h1']
    db = Db()
    if 'fileField1' in request.files:
        image = request.files['fileField1']
        if image.filename!="":
            dt = time.strftime("%Y%m%d-%H%M%S")
            image.save(staticpath + "staff_img\\" + dt + ".jpg")
            path = "/static/staff_img/" + dt + ".jpg"

            qry="UPDATE `staff` SET image='"+path+"' WHERE staff_id='"+str(id)+"'"
            res=db.update(qry)
            return '''<script>alert('updated');window.location='/staff_profile'</script>'''
#         else:
#             # dt = time.strftime("%Y%m%d-%H%M%S")
#             # image.save(staticpath + "college_img\\" + dt + ".jpg")
#             # path = "/static/staff_img/" + dt + ".jpg"
#             # qry = "UPDATE `staff` SET image='" + path + "' WHERE staff_id='" + str(id) + "'"
#             # res = db.update(qry)
            return '''<script>alert('hii');window.location='/staff_profile'</script>'''
    else:
        dt = time.strftime("%Y%m%d-%H%M%S")
        image.save(staticpath + "college_img\\" + dt + ".jpg")
        path = "/static/staff_img/" + dt + ".jpg"
        qry = "UPDATE `staff` SET image='" + path + "' WHERE staff_id='" + str(id) + "'"
        res = db.update(qry)
        return '''<script>alert('upUpdated');window.location='/staff_profile'</script>'''

@app.route('/view_staffbranch')
def view_staffbranch():
    db = Db()
    res = db.select("SELECT * FROM `branch`,`department` WHERE `branch`.`dept_id` = `department`.`dept_id` and `branch`.`dept_id` = '" +str(session['depid'])+ "'")
    return render_template('staff/view_staffbranch.html',data=res)

@app.route('/view_student1/<branch_id>')
def view_student1(branch_id):
    db = Db()
    session['branch_id'] = branch_id
    res=db.select("select student.*,department.dept_name,branch.branch_name FROM `department` INNER JOIN `student` ON `student`.`dept_id`=`department`.`dept_id` INNER JOIN `branch` ON `student`.`branch_id`= `branch`.`branch_id` WHERE student.`branch_id`='"+str(session['branch_id'])+"'" )
    print(res)
    return render_template('staff/view_student1.html',data=res)

@app.route('/view_student1_post', methods=['post'])
def view_student1_post():
    btn=request.form["button"]
    if btn== "Submit":
        s_year = request.form['select3']

        db = Db()
        qry="select student.*,department.dept_name,branch.branch_name FROM `department` INNER JOIN `student` ON `student`.`dept_id`=`department`.`dept_id` INNER JOIN `branch` ON `student`.`branch_id`= `branch`.`branch_id` WHERE student.`branch_id`='"+str(session['branch_id'])+"' AND `student`.`year` = '"+s_year+"' "
        # qry="SELECT `student`.*,`department`.`dept_name`,`branch`.`branch_name` FROM `student` INNER JOIN branch ON `student`.`branch_id`=`branch`.`branch_id` INNER JOIN `department` ON `branch`.`dept_id`=`department`.`dept_id` WHERE `student`.`college_id` ='"+str(session['lid'])+"' AND `student`.`year` = '"+year+"' "
        res = db.select(qry)

    elif btn=="view":
        s_year = request.form['select3']
        search = request.form['textbox10']
        db = Db()
        qry="select student.*,department.dept_name,branch.branch_name FROM `department` INNER JOIN `student` ON `student`.`dept_id`=`department`.`dept_id` INNER JOIN `branch` ON `student`.`branch_id`= `branch`.`branch_id` WHERE student.`branch_id`='"+str(session['branch_id'])+"' AND `student`.`s_name` LIKE '%" + search + "%' AND `student`.`year` LIKE '%" + s_year + "%'"
        # qry = "SELECT `student`.*,`department`.`dept_name`,`branch`.`branch_name` FROM `department` INNER JOIN `student` ON `student`.`dept_id`=`department`.`dept_id` INNER JOIN `branch` ON `student`.`branch_id`= `branch`.`branch_id` WHERE `student`.`s_name` LIKE '%" + search + "%'"
        res = db.select(qry)

    print("------------------------------")
    print(qry)
    return render_template('staff/view_student1.html',data=res)

@app.route('/delete_student1/<stud_id>')
def delete_student1(stud_id):
    db = Db()
    qry = "delete from student where stud_id='" + stud_id + "'"
    res = db.delete(qry)
    return '''<script>alert('Deleted');window.location='/view_student1'</script>'''

@app.route('/view_branches')
def view_branches():
    db = Db()
    res = db.select("SELECT * FROM `branch`,`department` WHERE `branch`.`dept_id` = `department`.`dept_id` and `branch`.`dept_id` = '" +str(session['depid'])+ "'")
    return render_template('staff/view_branches.html',data=res)


@app.route('/view_projects1/<branch_id>')
def view_projects1(branch_id):
    db=Db()
    session['branch_id'] = branch_id
    res=db.select("SELECT project.*,student.s_name,department.dept_name,branch.branch_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN department ON student.dept_id=department.dept_id INNER JOIN `branch` ON student.branch_id=branch.branch_id WHERE student.branch_id='" + str(session['branch_id']) + "' and project.status!='rejected'")

    return render_template('staff/view_projects1.html', data=res)
@app.route('/staff_approve_project/<pid>')
def staff_approve_project(pid):
    q="UPDATE `project` SET `status`='approved' WHERE `project_id`='"+pid+"'"
    d=Db()
    d.update(q)
    branch_id=session['branch_id']
    return view_projects1(branch_id)

@app.route('/staff_reject_project/<pid>')
def staff_reject_project(pid):
    q="UPDATE `project` SET `status`='rejected' WHERE `project_id`='"+pid+"'"
    d=Db()
    d.update(q)
    branch_id=session['branch_id']
    return view_projects1(branch_id)
@app.route('/view_projects1_post', methods=['post'])
def view_projects1_post():
    btn=request.form["button"]
    if(btn=="Submit"):
        year = request.form['select3']
        db = Db()
        # qry="select student.*,department.dept_name,branch.branch_name FROM `department` INNER JOIN `student` ON `student`.`dept_id`=`department`.`dept_id` INNER JOIN `branch` ON `student`.`branch_id`= `branch`.`branch_id` WHERE student.`branch_id`='"+str(session['branch_id'])+"' AND `student`.`s_year` = '"+s_year+"' "
        res = db.select("SELECT project.*,student.s_name,department.dept_name,branch.branch_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN department ON student.dept_id=department.dept_id INNER JOIN `branch` ON student.branch_id=branch.branch_id WHERE student.branch_id='" + str(session['branch_id']) + "' AND `project`.`p_year` = '"+year+"' and project.status!='rejected'")

    elif btn=="view":
        year = request.form['select3']
        search = request.form['textbox10']
        db = Db()
        # qry = "select student.*,department.dept_name,branch.branch_name FROM `department` INNER JOIN `student` ON `student`.`dept_id`=`department`.`dept_id` INNER JOIN `branch` ON `student`.`branch_id`= `branch`.`branch_id` WHERE student.`branch_id`='" + str(session['branch_id'])+"' and `student`.`s_name` LIKE '%" + search + "%'"
        res = db.select("SELECT project.*,student.s_name,department.dept_name,branch.branch_name FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN department ON student.dept_id=department.dept_id INNER JOIN `branch` ON student.branch_id=branch.branch_id WHERE student.branch_id='" + str(session['branch_id']) + "' AND `project`.`title` LIKE '%" + search + "%' or `project`.`project_area` LIKE '%" + search + "%' and `project`.`p_year` LIKE '%" + year + "%' and project.status!='rejected'")

    print("------------------------------")
    return render_template('staff/view_projects1.html',data=res)




@app.route('/view_newprojects1')
def view_newprojects1():
    db = Db()
    res = db.select("SELECT project.*,student.s_name,department.dept_name,branch.branch_name,college.c_name FROM student INNER JOIN project ON project.stud_id=student.stud_id INNER JOIN department ON student.dept_id=department.dept_id INNER JOIN branch ON student.branch_id=branch.branch_id INNER JOIN college ON student.college_id=college.college_id  WHERE college.login_id='" + str(session['lid']) + "'")
    print(res)
    return render_template('staff/view_newprojects1.html', data=res)


@app.route('/view_newprojects1_post', methods=['post'])
def view_newprojects1_post():
    search = request.form['search_project']
    db = Db()
    res = db.select("SELECT project.*,student.s_name,department.dept_name,branch.branch_name,college.c_name FROM student INNER JOIN project ON project.stud_id=student.stud_id INNER JOIN department ON student.dept_id=department.dept_id INNER JOIN branch ON student.branch_id=branch.branch_id INNER JOIN college ON student.college_id=college.college_id  where `project`.`title` LIKE '%" + search + "%' or `project`.`year` LIKE '%" + search + "%' or `student`.`s_name` LIKE '%" + search + "%' or `department`.`dept_name` LIKE '%" + search + "%' or `branch`.`branch_name` LIKE '%" + search + "%' or `college`.`c_name` LIKE '%" + search + "%'")
    print(res)
    return render_template('admin/view_newprojects1.html', data=res)

@app.route('/reg_students')
def reg_students():
    db = Db()
    res = db.select("SELECT * FROM `department` WHERE college_id = '" + str(session["lid"]) + "'")
    print(res)
    re = db.select("SELECT * FROM `branch` WHERE college_id = '" + str(session["lid"]) + "'")
    print(re)
    return render_template('college/reg_students.html', dept=res, branch=re)

@app.route('/select_dept2')
def select_dept2():
    db=Db()
    res = db.select("SELECT * FROM `department` WHERE `college_id` ='" + str(session['lid']) + "'")
    print(res)
    return render_template('college/reg_students.html', dept=res)

@app.route('/brnch2_ajax',methods=["post"])
def brnch2_ajax():
    depid=request.form['depid']
    db=Db()
    q="SELECT * FROM `branch` WHERE `dept_id`='"+depid+"'"
    res=db.select(q)
    print(res)
    print(q)

    return jsonify(status="ok",data=res)

@app.route('/reg_students_post',methods=['post'])
def reg_students_post():
    db = Db()
#     # name = request.form['textfield']
#     # email = request.form['textfield2']
#     # phone = request.form['textfield3']
    department = request.form['select']
    # print("hi")
    branch = request.form['select2']
    # print("hii")
    year=request.form['select3']
    # print("hiii")
    f = request.files['fileField']
    dt = time.strftime("%Y%m%d-%H%M%S")
    f.save(staticpath + "student\\" + dt + ".csv")
    path = "/static/student/" + dt + ".csv"

    p = pd.read_csv(staticpath + "student\\" + dt + ".csv")
    print(p)
    print(len(p))
    for i in range(len(p)):
        student_name=str(p["s_name"][i])
        email=str(p["email"][i])
        phoneno=str(p["phone"][i])
        import random
        passw = str(random.randint(0000, 9999))
    # print("Column headings:")
    # print(p["s_name"])
        lid=db.insert("insert into login(username,password,user_type) VALUES('" + email + "','" + passw + "','student')")
        db.insert("insert into student (college_id,`s_name`,email,phone,dept_id,branch_id,`year`,login_id) VALUES ('"+str(session['lid'])+"','" + student_name + "','" + email + "','" + phoneno + "','" + department + "','" + branch + "','" + year + "','"+str(lid)+"')")
    return '''<script>alert('students registered successfully');window.location='/reg_students'</script>'''




@app.route('/reg_staffs')
def reg_staffs():
    db = Db()
    res = db.select("SELECT * FROM `department` WHERE college_id = '" + str(session["lid"]) + "'")
    print(res)
    return render_template('college/reg_staffs.html', dept=res)
@app.route('/reg_staffs_post',methods=['post'])
def reg_staffs_post():
    db = Db()
    import pandas as pd
    department = request.form['select']

    f = request.files['fileField']
    dt = time.strftime("%Y%m%d-%H%M%S")
    f.save(staticpath + "staff\\" + dt + ".csv")
    path = "/static/staff/" + dt + ".csv"
    p = pd.read_csv(staticpath + "staff\\" + dt + ".csv")
    print(p)
    print(len(p))
    for i in range(len(p)):
        staff_name=str(p["name"][i])
        email=str(p["email"][i])
        phoneno=str(p["phone"][i])
        import random
        passw = str(random.randint(0000, 9999))

        lid = db.insert("insert into login(username,password,user_type) VALUES('" + email + "','" + passw + "','staff')")
        db.insert("insert into staff (college_id,`name`,email,phone,dept_id,login_id) VALUES ('" + str(session['lid']) + "','" + staff_name + "','" + email + "','" + phoneno + "','" + department + "','" + str(lid) + "')")
    return '''<script>alert('staffs added successfully');window.location='/reg_staffs'</script>'''

# @app.route('/studhome')
# def studhome():
#     return render_template('student/studhome.html')
#
# @app.route('/stud_profile')
# def stud_profile():
#     db = Db()
#     qry="SELECT student.*,department.dept_name,branch.branch_name FROM student INNER JOIN department on student.dept_id=department.dept_id INNER JOIN branch on student.branch_id=branch.branch_id WHERE student.login_id ='" +str(session['lid'])+ "'"
#     # qry = "SELECT * FROM `student`,department,branch WHERE `login_id` ='" +str(session['lid'])+ "'"
#     res = db.selectOne(qry)
#     return render_template('student/stud_profile.html', data=res)
#
# @app.route('/stud_changepwd')
# def stud_changepwd():
#     return render_template('student/stud_changepwd.html')
#
# @app.route('/stud_changepwd_post',methods=["post"])
# def stud_changepwd_post():
#     db = Db()
#     newpwd = request.form['textfield2']
#     confnewpwd = request.form['textfield3']
#
#     if newpwd==confnewpwd:
#         qry = "UPDATE `login` SET password='" + newpwd + "' where login_id ='" +str(session["lid"])+ "'"
#         res = db.update(qry)
#         return '''<script>alert('updated');window.location='/index'</script>'''
#     else:
#         return '''<script>alert('Password mismatch');window.location='/stud_changepwd'</script>'''
#
# @app.route('/pro_upload')
# def pro_upload():
#     return render_template('student/pro_upload.html')
#
# @app.route('/pro_upload_post',methods=["post"])
# def pro_upload_post():
#     lid=request.form['login_id']
#     title=request.form['pname']
#     abstract=request.form['ab']
#     demo = request.form['demo']
#     report=request.form['rep']
#     area=request.form['area']
#     githublink=request.form['git']
#     year = request.form['year']
#     timestr_report = time.strftime("%Y%m%d-%H%M%S")
#     a = base64.b64decode(report)
#     fh = open(staticpath + "report\\" + timestr_report + ".pdf", "wb")
#     reportpath = "/static/report/" + timestr_report + ".pdf"
#     # fh.write(a)
#     # fh.close()
#     timestr = time.strftime("%Y%m%d-%H%M%S")
#     print(timestr)
#     a = base64.b64decode(demo)
#     fh = open(staticpath + "Demo\\" + timestr + ".mp4", "wb")
#     demopath = "/static/Demo/" + timestr + ".mp4"
#     # fh.write(a)
#     # fh.close()
#     print("githublink")
#     journal=request.form['j1']
#     timestr = time.strftime("%Y%m%d-%H%M%S")
#     print(timestr)
#     a = base64.b64decode(journal)
#     fh = open(staticpath + "journal\\" + timestr + ".pdf", "wb")
#     jpath = "/static/journal/" + timestr + ".pdf"
#     fh.write(a)
#     fh.close()
#     print("journal")
#     q="insert into project(title,abstract,report,journal,`p_year`,demo,githublink,stud_id,project_area) VALUES ('" +title+ "','"+abstract+"','"+reportpath+"','"+jpath+"','"+year+"','"+demopath+"','"+githublink+"','"+lid+"','"+area+"')"
#     print(q)
#     # pid=c.insert(q)
#     return '''<script>alert('Project Uploaded Successfully');window.location='/pro_upload'</script>'''
#
# @app.route('/my_projects')
# def my_projects():
#     db = Db()
#     q="SELECT project.* FROM project WHERE project.stud_id='" +str(session['lid'])+ "'"
#     res = db.selectOne(q)
#     return render_template('student/my_projects.html', data=res)

@app.route('/user_add_rating/<project_id>')
def user_add_rating(project_id):
    session["pid_id"]=project_id
    return render_template('viewmoreproject.html')

@app.route('/user_add_rating_post')
def user_add_rating_post():
    rate=request.args.get("rat")
    print(rate,"--------")
    db=Db()
    qry="select * from rating where gustlid='"+str(session['lid'])+"' and pid='"+str(session['pid_id'])+"'"
    res=db.select(qry)
    if len(res)==0:
        q = db.insert("insert into `rating`(`pid`,`gustlid`,`rate`,`date`)VALUES ('"+str(session['pid_id'])+"','"+str(session['lid'])+"','" + rate + "',curdate())")
        print(q)
    else:
        q = db.insert("update `rating` set rate='"+rate+"'  where gustlid='"+str(session['lid'])+"' and pid='"+str(session['pid_id'])+"'")
        print(q)
    return '''<script>alert('rating added');window.location='/user_add_rating'</script>'''
#=-------------------------------Android------------------------------------------------------

@app.route('/andlogin_post', methods=['post'])
def andlogin_post():
   db=Db()
   uname = request.form['uname']
   password = request.form['passw']
   q="select * from login WHERE username='"+uname+"'and password='"+password+"'"
   ss=db.selectOne(q)
   print(q)
   if ss is not None :
       return jsonify(status="ok",lid=ss["login_id"],type=ss['user_type'])
   else:
        return jsonify(status="not ok")

@app.route('/view_studentprofile_postt',methods=["post"])
def view_studentprofile_postt():
    print("iiiiiiiiiiiiiiiiiiiiiiiii")
    lid=request.form['lid']
    db = Db()
    q="SELECT student.*,branch.branch_name,department.dept_name FROM student INNER JOIN branch ON student.branch_id=branch.branch_id INNER JOIN department ON branch.dept_id=department.dept_id WHERE student.login_id='"+str(lid)+"'"
    print(q)
    ss = db.selectOne(q)
    print(ss)
    if ss is not None :
       return jsonify(status="ok",name=ss['s_name'],email=ss['email'],phone=ss['phone'],department=ss['dept_name'],branch=ss['branch_name'],image=ss['image'])
    else:
        return jsonify(status="not ok")

@app.route('/view_myprojects',methods=["post"])
def view_studentproject_post():
    lid=request.form['login_id']
    db = Db()
    q="SELECT project.* FROM project WHERE project.stud_id='"+str(lid)+"'"
    print(q)
    ss = db.select(q)
    print(ss)
    if ss is not None:
        return jsonify(status="ok",users=ss)
    else:
        return jsonify(status="not ok")
@app.route('/projectupload_post',methods=["post"])
def projectupload_post():


    distincts=[]


    lid=request.form['login_id']
    print("lid=====+",lid);
    db=Db()
    title=request.form['pname']
    print("title====+",title);
    abstract=request.form['ab']
    # print("abstract=====+",abstract);
    demo = request.form['demo']
    # print("demo====+",demo);
    report=request.form['rep']
    # print("rep=========+",report);
    area=request.form['area']
    githublink=request.form['git']
    print("githublink====+",githublink);



    print("abstract");
    timestr_report = time.strftime("%Y%m%d-%H%M%S")
    print(timestr_report)
    a = base64.b64decode(report)
    fh = open(staticpath + "report\\" + timestr_report + ".pdf", "wb")
    reportpath = "/static/report/" + timestr_report + ".pdf"
    fh.write(a)
    fh.close()
    print("report")
    import PyPDF2






    year=request.form['year']
    print("year")

    timestr = time.strftime("%Y%m%d-%H%M%S")
    print(timestr)
    a = base64.b64decode(demo)
    fh = open(staticpath + "Demo\\" + timestr + ".mp4", "wb")
    demopath = "/static/Demo/" + timestr + ".mp4"
    fh.write(a)
    fh.close()
    print("888888888888888888888888888888888888888888888888888888888888888888888888888888888888888888888888888888888888")
    print("githublink")
    journal=request.form['j1']
    timestr = time.strftime("%Y%m%d-%H%M%S")
    print(timestr)
    a = base64.b64decode(journal)
    fh = open(staticpath + "journal\\" + timestr + ".pdf", "wb")
    jpath = "/static/journal/" + timestr + ".pdf"
    fh.write(a)
    fh.close()
    print("journal")
    c=Db()
    q="insert into project(title,abstract,report,journal,`p_year`,demo,githublink,stud_id,project_area) VALUES ('" +title+ "','"+abstract+"','"+reportpath+"','"+jpath+"','"+year+"','"+demopath+"','"+githublink+"','"+lid+"','"+area+"')"
    print(q)
    pid=c.insert(q)


    qq="INSERT INTO `rating`(`pid`,`gustlid`,DATE,rate)VALUES('"+str(pid)+"','1',CURDATE(),'0')"
    c.insert(qq)



    doc = aw.Document("C:\\Users\\User\\PycharmProjects\\project_management12\\static\\report\\" + timestr_report + ".pdf")

    # convert PDF to Word DOCX format
    doc.save("C:\\Users\\User\\PycharmProjects\\project_management12\\static\\convert_doc\\"+timestr_report + ".docx")
    import docx
    doc = docx.Document("C:\\Users\\User\\PycharmProjects\\project_management12\\static\\convert_doc\\"+timestr_report + ".docx")  # Creating word reader object.
    data = ""
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
        data = '\n'.join(fullText)

    print(data)

    data.replace(",", "").replace(".", "").replace("0", "").replace("1", "").replace("2", "").replace("3",
                                                                                                    "").replace(
        "4", "").replace("5", "").replace("6", "").replace("7", "").replace("8", "").replace("0", "").replace("!",
                                                                                                              "").replace("@", "").replace("'","")
    ss = word_tokenize(data)
    print(ss)
    filtered_sentence = []

    for w in ss:
        if w not in stop_words:
            filtered_sentence.append(w)

    print(ss)
    print(filtered_sentence)
    # aa = list(session["wordslist"])
    # print(aa)
    aa = []
    for i in filtered_sentence:
        aa.append(i)
    session["wordslist"] = aa
    print(aa)
    dd=[]


    from collections import Counter

    counts = Counter(aa)
    print(counts)
    k = counts.keys()
    print(k)
    v = counts.values()

    # for i in range(0, len(k)):
    #     print(list(k)[i])
    #
    #     print(list(v)[i])
    for p in aa:
        try:
            pp=str(p).lower().replace("0","").replace("1","").replace("2","").replace("3","").replace("4","").replace("5","").replace("6","").replace("7","").replace("8","").replace("9","").replace(".","").replace("-","").replace(",","").replace(" ","").replace("/","").replace("\\","")
            if pp not in gl_list:
                print(" ")
                print(pp)
                pp=pp.lstrip().rstrip().strip()
                if pp != "":
                    # qry="SELECT * FROM `key_words` WHERE `key`='"+str(pp)+"' AND `pid`='"+str(pid)+"'"
                    # print(qry)
                    # fata=db.select(qry)
                    # print(fata)
                    # if len(fata)==0:
                    #     print("(((((((((((((((((((((((((")

                        if pp not in distincts:

                            q = "INSERT INTO `key_words`(`key`,`count`,`pid`)VALUES('" + str(pp) + "','1','" + str(pid) + "')"
                            d = Db()
                            print(q)
                            d.insert(q)

                            distincts.append(pp)


                        else:

                            print("else")

                            q = "update `key_words`set `count`=`count`+1 where `key`='"+str(pp)+"' and pid='"+str(pid)+"'"
                            d = Db()
                            print(q)
                            d.update(q)

        except :
            pass

    try:
        import prt
        prt.abcde(staticpath + "journal\\" + timestr + ".pdf", str(pid))
    except:
        pass
    return jsonify(status="ok")

@app.route('/and_delete_my_project',methods=["post"])
def and_delete_my_project():
    pid=request.form["pid"]
    q="DELETE FROM `project` WHERE `project_id`='"+pid+"'"
    d=Db()
    d.delete(q)
    return jsonify(status="ok")

@app.route('/and_view_department',methods=["post"])
def and_view_department():
    id=request.form["id"]
    print(id)
    q="SELECT department.*,college.*,student.*,department.dept_id AS depidd FROM department,college ,student WHERE department.college_id=college.college_id AND student.college_id=department.college_id AND student.login_id='"+str(id)+"' "
    # q="DELETE FROM `project` WHERE `project_id`='"+pid+"'"
    db=Db()
    ss=db.select(q)
    print(ss)
    return jsonify(status="ok",data=ss)



@app.route('/andview_projectguest1',methods=['post'])
def andview_projectguest1():
    did=request.form["did"]
    print(did)
    db=Db()
    q="SELECT project.*,student.s_name,college.c_name,project.p_year FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id where  `student`.`dept_id`='"+did+"' and  project.`status`='approved'"
    ss=db.select(q)
    print(q)
    print("ghosttttttttttttttttttttttttt",len(ss))
    print(ss)
    return jsonify(status="ok",data = ss)


@app.route('/andview_studnetproject_search',methods=['post'])
def andview_studnetproject_search():
    search = request.form['search_project']
    dept_id = request.form['dept_id']
    db = Db()

    res=search.split(" ")
    print(res)
    # ss = db.select(" SELECT project.project_id,project.title,student.s_name,college.c_name,project.p_year FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id where project.`title` LIKE '%" + search + "%' or project.`p_year`LIKE '%" + search + "%'")
    # ss = db.select(" SELECT project.project_id,project.title,student.s_name,college.c_name,project.p_year FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id where project.`title` LIKE '%" + search + "%' or project.`p_year`LIKE '%" + search + "%'")

    pid=[]
    ls=[]
    for i in res:
        if i!='':
            Q="SELECT distinct `key_words`.`pid`,key_words.count, project.project_id,project.title,student.s_name,college.c_name,project.p_year FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN `key_words`  ON `project`.`project_id`=`key_words`.`pid` WHERE  student.dept_id='"+dept_id+"' and (`key_words`.`key` LIKE '%" + i + "%'  OR( `project`.title LIKE '%"+i+"%')OR( `project`.project_area LIKE '%"+i+"%'))   ORDER BY `count` DESC"
            ss=db.select(Q)
            print(Q)
            for i in ss:

                if i["project_id"] in pid:
                    print("---")
                    pass
                else:
                    pid.append(i["project_id"])
                    a={'title':i["title"],'s_name':i["s_name"],'c_name':i["c_name"],'p_year':i["p_year"],'project_id':i["project_id"]}
                    print("---",a)
                    ls.append(a)



    print(ls)
    print(len(ls))
    return jsonify(status="ok",data=ls)



@app.route('/andview_studnetproject_search_year',methods=['post'])
def andview_studnetproject_search_year():
    year = request.form['search_project']
    dept_id = request.form['dept_id']
    db = Db()
    ss="SELECT * FROM `project` INNER JOIN `student` ON `student`.`login_id`=`project`.`stud_id` INNER JOIN `college` ON `college`.login_id=`student`.`college_id`  WHERE `p_year`='"+year+"' AND `student`.`dept_id`='"+dept_id+"' and  project.`status`='approved'"
    print(ss,"--------------")
    ls=db.select(ss)
    print(ls)
    return jsonify(status="ok",data=ls)

@app.route('/andcoustomview_projectguest',methods=["post"])
def andcustomview_projectguest():
    db=Db()
    q="select title,s_name from project inner join student on student.stud_id=project.stud_id"
    print(q)
    ss=db.select(q)
    print(ss)
    return jsonify(status="ok")

@app.route('/projectdetails',methods=["post"])
def projectdetails():
    pid=request.form['pid']
    print(pid)
    db=Db()
    ss=db.selectOne("select title,abstract,journal,demo,report,githublink from project where project_id='"+str(pid)+"'");
    print(ss)
    return jsonify(status="ok",title = ss['title'],abst = ss['abstract'],report=ss['report'],demo=ss['demo'],journal=ss['journal'],githublink = ss['githublink'])

@app.route('/stud_view_rating',methods=["post"])
def stud_view_rating():
    pid=request.form["pid"]
    q="SELECT rating.*,`gust`.* FROM `rating` INNER JOIN `gust` ON `rating`.`gustlid`=`gust`.`glid` WHERE `rating`.`pid`='"+pid+"'"
    ss = "SELECT avg(rate) as a FROM rating WHERE `rating`.`pid`='"+ pid+"'"
    d=Db()
    res=d.select(q)
    r = d.selectOne(ss)
    return jsonify(status="ok",users=res,avg=r["a"])


# @app.route('/edit_profile',methods=["post"])
# def edit_profile():
#     db=Db()
#     name= request.form['edit_name1']
#     email = request.form['edit_email1']
#     phone = request.form['edit_phone1']
#     department = request.form['edit_department1']
#     branch = request.form['edit_branch1']
#     ss="UPDATE student SET s_name='"+name+"',email='"+email+"',phone='"+phone+"',department='"+department+"',branch='"+branch+"'"
@app.route('/andview_department',methods=["post"])
def andview_department():
    lid=request.form['login_id']
    db=Db()
    ss="SELECT dept_name FROM department INNER JOIN college ON department.college_id=college.college_id INNER JOIN student ON student.college_id=college.college_id where  student.login_id='"+str(lid)+"'"

# @app.route('/andsearchproject')

#--------------------------Gust___________________
@app.route('/gust_reg',methods=["post"])
def gust_reg():
    name=request.form["name"]
    email=request.form["email"]
    passw=request.form["passw"]
    d=Db()
    q="INSERT INTO `login`(`username`,`password`,`user_type`)VALUES('"+email+"','"+passw+"','gust')"
    lid=d.insert(q)
    q1="INSERT INTO `gust`(`name`,`email`,`glid`)VALUES('"+name+"','"+email+"','"+str(lid)+"')"
    d.insert(q1)
    return jsonify(status="ok")
@app.route('/gust_view_all_project',methods=["post"])
def gust_view_all_project():
    lid=request.form["lid"]
    d=Db()
    q="SELECT `project`.*,`student`.*,`department`.*,`college`.*,AVG(rate) as r FROM rating INNER JOIN project ON `project`.project_id=rating.pid INNER JOIN `student` ON `project`.`stud_id`=`student`.`login_id` INNER JOIN `department` ON `student`.`dept_id`=`department`.`dept_id` INNER JOIN `college` ON `department`.`college_id`=`college`.`college_id` WHERE project.`status`= 'approved' GROUP BY pid"
    res=d.select(q)
    print(res)
    ls=[]
    if len(res)>0:

        for i in res:
            q="SELECT * FROM `payment` WHERE `gustlid`='"+lid+"' AND `projectid`='"+str(i["project_id"])+"'"
            c=Db()
            respay=c.selectOne(q)
            if respay is not None:
                a={'title':i['title'],'abstract':i['abstract'],'project_id':i['project_id'],'s_name':i['s_name'],'c_name':i['dept_name'],'p_year':i['year'],'status':'yes','rat':i['r']}
                ls.append(a)
            else:
                a = {'title': i['title'], 'abstract': i['abstract'], 'project_id': i['project_id'], 's_name': i['s_name'],
                     'c_name': i['dept_name'], 'p_year': i['year'], 'status': 'no','rat':i['r']}
                ls.append(a)

    return jsonify(status="ok",data=ls)

@app.route('/user_payment',methods=['post'])
def user_payment():
    gustlid=request.form["lid"]
    projectid=request.form["pid"]
    bank=request.form["bank"]
    pin=request.form["pin"]
    amount="1000"
    d=Db()
    bank="SELECT `amount` FROM `bank` WHERE `accpin`='"+pin+"' AND `bank`='"+bank+"'"
    bankres=d.select(bank)
    ls=[]
    if bankres is not None:
        if int(amount)<=int(bankres["amount"]):
            q="INSERT INTO `payment`(`gustlid`,`projectid`,`date`,`amount`)VALUES('"+gustlid+"','"+projectid+"',curdate())"
            dd=Db()
            dd.insert(q)

            qq="UPDATE `bank` SET `amount`=`amount`-'"+amount+"' WHERE `accpin`='"+pin+"' AND `bank`='"+bank+"'"
            dd.insert(qq)
        else:
            return jsonify(status="low")
    else:
        return jsonify(status="no")

@app.route('/gust_send_rating',methods=["post"])
def gust_send_rating():
    pid=request.form["pid"]
    glid=request.form["lid"]
    d=Db()
    rate=request.form["rate"]
    qry="select * from rating where pid='"+pid+"' and gustlid='"+glid +"' "
    res = d.selectOne(qry)
    print(qry)
    if res is not None:
        w="update rating set rate ='"+rate+"' where pid='"+pid+"' and gustlid='"+glid+"'"
        print(w)
        d = Db()
        d.update(w)
    else:
        q="INSERT INTO `rating`(`pid`,`gustlid`,`date`,`rate`)VALUES('"+pid+"','"+glid+"',curdate(),'"+rate+"')"
        print(q)
        d=Db()
        d.insert(q)
    return jsonify(status="ok")

@app.route('/gust_inset_fav',methods=["post"])
def gust_inset_fav():
    pid=request.form["pid"]
    glid=request.form["lid"]
    # rate=request.form["rate"]

    q="INSERT INTO `favorite`(`pid`,`glid`,`date`)VALUES('"+pid+"','"+glid+"',curdate())"
    d=Db()
    d.insert(q)
    return jsonify(status="ok")
@app.route('/gust_delete_fav',methods=["post"])
def gust_delete_fav():
    favid=request.form["fid"]
    q="DELETE FROM `favorite` WHERE `fid`='"+favid+"'"
    d=Db()
    d.delete(q)
    return jsonify(status="ok")
@app.route('/Viewfavourite',methods=["post"])
def Viewfavourite():
    lid=request.form["lid"]
    q="SELECT `project`.*,`favorite`.* ,`student`.*,`department`.*,`college`.* FROM `project` INNER JOIN `favorite` ON `project`.`project_id`=`favorite`.`pid`  INNER JOIN `student` ON `project`.`stud_id`=`student`.`login_id` INNER JOIN `department` ON `student`.`dept_id`=`department`.`dept_id` INNER JOIN `college` ON `department`.`college_id`=`college`.`college_id` WHERE `glid`='"+lid+"'"
    d=Db()
    res=d.select(q)
    return jsonify(status="ok",data=res)



@app.route('/andview_gust_project_search',methods=['post'])
def andview_gust_project_search():
    lid=request.form["lid"]
    search = request.form['search_project']
    db = Db()
    # ==========================

    restt = search.split(" ")
    print(restt)
    lss = []
    ls = []
    pid = []
    qry = "SELECT project.project_id,project.title,student.s_name,college.c_name,project.p_year,abstract,avg(rate) as rat FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN `ngram` ON `ngram`.`pid`=`project`.`project_id`INNER JOIN `department` ON department.dept_id=student.dept_id inner join rating on `rating`.`pid`=`project`.`project_id` WHERE ngram.`word` = '" + search + "' or project.`title` = '" + search + "' GROUP BY  project.project_id,project.title,student.s_name,college.c_name,project.p_year,abstract"
    db = Db()
    res = db.select(qry)
    print(qry)
    print(res)
    print(len(res), "------")
    if len(res) > 0:
        print(res)
        print(len(res))
        for i in res:
            print(i,"iiiiiiiiiiiiii")
            if i["project_id"] in pid:
                print("---", i["project_id"])
                pass
            else:
                print(pid, "--------------")
                pid.append(i["project_id"])
                q = "SELECT * FROM `payment` WHERE `gustlid`='" + lid + "' AND `projectid`='" + str(
                    i["project_id"]) + "'"
                c = Db()
                respay = c.selectOne(q)
                if respay is not None :

                    a = {'title': i['title'], 'abstract': i['abstract'], 'project_id': i['project_id'],
                         's_name': i['s_name'],
                         'c_name': i['c_name'], 'p_year': i['p_year'], 'status': 'yes','rat':i["rat"]}
                    ls.append(a)
                else:
                    a = {'title': i['title'], 'abstract': i['abstract'], 'project_id': i['project_id'],
                         's_name': i['s_name'],
                         'c_name': i['c_name'], 'p_year': i['p_year'], 'status': 'no','rat':i["rat"]}
                    ls.append(a)

        return jsonify(status="ok", data=ls)

    else:
        print("eeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeee")
        # ss = db.select(" SELECT project.project_id,project.title,student.s_name,college.c_name,project.p_year FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id where project.`title` LIKE '%" + search + "%' or project.`p_year`LIKE '%" + search + "%'")
        ss = db.select(
            " SELECT project.project_id,project.title,student.s_name,college.c_name,project.p_year,abstract,avg(rate) as rat  FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id inner join rating on `rating`.`pid`=`project`.`project_id` GROUP BY  project.project_id,project.title,student.s_name,college.c_name,project.p_year,abstract")
        print(ss)
        pid = []
        ls = []
        for i in ss:
            p = 0
            for j in restt:
                if j != '':
                    Q = "SELECT abstract, project.project_id,project.title,student.s_name,college.c_name,project.p_year FROM project INNER JOIN student ON project.stud_id=student.login_id INNER JOIN college ON student.college_id=college.college_id INNER JOIN `key_words`  ON `project`.`project_id`=`key_words`.`pid` where `key_words`.`pid`='" + str(
                        i["project_id"]) + "' and key_words.`key` ='" + str(j) + "'"
                    print(Q)
                    pp = db.selectOne(Q)
                    print("?*************************************************")
                    if pp is not None:

                        if i["project_id"] in pid:
                            print("---", i["project_id"])
                            pass
                        else:
                            print(pid, "--------------")
                            pid.append(i["project_id"])
                            q = "SELECT * FROM `payment` WHERE `gustlid`='" + lid + "' AND `projectid`='" + str(
                                i["project_id"]) + "'"
                            c = Db()
                            respay = c.selectOne(q)
                            if respay is not None:
                                a = {'title': i['title'], 'abstract': i['abstract'], 'project_id': i['project_id'],
                                     's_name': i['s_name'],
                                     'c_name': i['c_name'], 'p_year': i['p_year'], 'status': 'yes','rat':i["rat"]}
                                ls.append(a)
                            else:
                                a = {'title': i['title'], 'abstract': i['abstract'], 'project_id': i['project_id'],
                                     's_name': i['s_name'],
                                     'c_name': i['c_name'], 'p_year': i['p_year'], 'status': 'no','rat':i["rat"]}
                                ls.append(a)
            lss.append(p)
        pid=[]


        print(ls)
        return jsonify(status="ok", data=ls)

















@app.route('/razorpay',methods=['post'])
def razorpay():
    phone_number=request.form['phone1']
    email_id=request.form['email1']
    refference_id=request.form['refer']

    gustlid = request.form["lid"]
    projectid = request.form["pid"]

    import requests
    import json

    url = "https://api.razorpay.com/v1/payment_links"

    payload = json.dumps({
        "amount": 100,
        "currency": "INR",
        "accept_partial": True,
        "first_min_partial_amount": 100,
        "expire_by": 1691097057,
        "reference_id": refference_id,
        "description": "Payment for policy no #23456",
        "customer": {
            "name": "Angel",
            "contact": phone_number,
            "email": email_id
        },
        "notify": {
            "sms": True,
            "email": True
        },
        "reminder_enable": True,
        "notes": {
            "policy_name": "Jeevan Bima"
        },
        "callback_url": "https://example-callback-url.com/",
        "callback_method": "get"
    })
    headers = {
        'Authorization': 'Basic cnpwX3Rlc3RfNTI4NkdHT1I4dUhwdUI6b01GR3kzYjF3WVNEWElSclBJMU05Y29L',
        'Content-Type': 'application/json'
    }

    response = requests.request("POST", url, headers=headers, data=payload)

    print(response.text)

    q = "INSERT INTO `payment`(`gustlid`,`projectid`,`date`,`amount`)VALUES('" + gustlid + "','" + projectid + "',curdate(),'1000')"
    dd = Db()
    dd.insert(q)
    return jsonify(status="ok")
@app.route('/andview_department12',methods=['post'])
def andview_department12():
    db=Db()
    ss=db.select("SELECT * from department")
    print(ss)
    return jsonify(status="ok",data = ss)



if __name__ == '__main__':
    app.run(debug=True,host="0.0.0.0")





